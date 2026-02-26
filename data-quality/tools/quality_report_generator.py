"""
quality_report_generator.py
---------------------------
Genera informes de cobertura de calidad del dato en PDF, DOCX o Markdown.

Uso:
    python tools/quality_report_generator.py --format pdf --output output/report.pdf --input-json '{"title": ...}'
    python tools/quality_report_generator.py --format docx --output output/report.docx --input-file data.json
    python tools/quality_report_generator.py --format md --output output/report.md --input-json '{"title": ...}'

Formato del JSON de entrada:
    {
        "title": "Informe de Cobertura de Calidad del Dato",
        "domain": "semantic_financial",
        "scope": "Dominio completo",
        "generated_at": "2025-02-26",
        "summary": {
            "tables_analyzed": 5,
            "rules_total": 12,
            "rules_ok": 8,
            "rules_ko": 2,
            "rules_warning": 1,
            "rules_not_executed": 1,
            "coverage_estimate": "62%",
            "gaps_critical": 4,
            "gaps_moderate": 3,
            "gaps_low": 2,
            "rules_created_this_session": 3
        },
        "tables": [
            {
                "name": "account",
                "coverage_estimate": "75%",
                "completeness": "OK",
                "uniqueness": "Gap",
                "validity": "Parcial",
                "consistency": "N/A",
                "rules": [
                    {
                        "name": "dq-account-completeness-id",
                        "dimension": "completeness",
                        "status": "OK",
                        "pass_pct": 100.0,
                        "description": "El campo id nunca es nulo"
                    }
                ],
                "gaps": [
                    {
                        "column": "account_id",
                        "dimension": "uniqueness",
                        "priority": "CRITICO",
                        "description": "Clave primaria sin verificacion de unicidad"
                    }
                ]
            }
        ],
        "rules_created": [
            {
                "name": "dq-account-uniqueness-id",
                "table": "account",
                "dimension": "uniqueness",
                "status": "created"
            }
        ],
        "recommendations": [
            "Resolver las 2 reglas en estado KO antes de crear nuevas",
            "Cubrir los gaps criticos de claves primarias con uniqueness"
        ]
    }
"""

import argparse
import json
import sys
import os
from datetime import date


# ---------------------------------------------------------------------------
# Helpers comunes
# ---------------------------------------------------------------------------

STATUS_ICONS = {
    "OK": "[OK]",
    "KO": "[KO]",
    "WARNING": "[WARN]",
    "N/A": "[N/A]",
    "Gap": "[GAP]",
    "Parcial": "[PARCIAL]",
    "created": "[CREADA]",
    "error": "[ERROR]",
}

PRIORITY_LABELS = {"CRITICO": "CRITICO", "ALTO": "ALTO", "MEDIO": "MEDIO", "BAJO": "BAJO"}


def load_input(input_json: str | None, input_file: str | None) -> dict:
    if input_json:
        return json.loads(input_json)
    if input_file:
        with open(input_file, encoding="utf-8") as f:
            return json.load(f)
    # Intentar leer de stdin
    raw = sys.stdin.read().strip()
    if raw:
        return json.loads(raw)
    raise ValueError("No se proporcionaron datos de entrada (--input-json, --input-file o stdin)")


def ensure_output_dir(output_path: str) -> None:
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)


# ---------------------------------------------------------------------------
# Generacion de Markdown
# ---------------------------------------------------------------------------

def build_markdown(data: dict) -> str:
    s = data.get("summary", {})
    lines = []

    # Cabecera
    lines.append(f"# {data.get('title', 'Informe de Cobertura de Calidad del Dato')}")
    lines.append("")
    lines.append(f"**Dominio / Coleccion**: {data.get('domain', '-')}")
    lines.append(f"**Scope**: {data.get('scope', '-')}")
    lines.append(f"**Fecha de generacion**: {data.get('generated_at', str(date.today()))}")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Resumen ejecutivo
    lines.append("## Resumen Ejecutivo")
    lines.append("")
    lines.append(f"| Metrica | Valor |")
    lines.append(f"|---------|-------|")
    lines.append(f"| Tablas analizadas | {s.get('tables_analyzed', '-')} |")
    lines.append(f"| Reglas existentes | {s.get('rules_total', '-')} |")
    lines.append(f"| Estado OK | {s.get('rules_ok', '-')} |")
    lines.append(f"| Estado KO | {s.get('rules_ko', '-')} |")
    lines.append(f"| Estado WARNING | {s.get('rules_warning', '-')} |")
    lines.append(f"| Sin ejecutar | {s.get('rules_not_executed', '-')} |")
    lines.append(f"| Cobertura estimada | **{s.get('coverage_estimate', '-')}** |")
    lines.append(f"| Gaps criticos | {s.get('gaps_critical', '-')} |")
    lines.append(f"| Gaps moderados | {s.get('gaps_moderate', '-')} |")
    lines.append(f"| Gaps bajos | {s.get('gaps_low', '-')} |")
    if s.get("rules_created_this_session"):
        lines.append(f"| Reglas creadas esta sesion | {s['rules_created_this_session']} |")
    lines.append("")

    # Cobertura por tabla
    tables = data.get("tables", [])
    if tables:
        lines.append("## Cobertura por Tabla")
        lines.append("")
        lines.append("| Tabla | Completeness | Uniqueness | Validity | Consistency | Cobertura |")
        lines.append("|-------|-------------|------------|----------|-------------|-----------|")
        for t in tables:
            lines.append(
                f"| {t['name']} "
                f"| {t.get('completeness', '-')} "
                f"| {t.get('uniqueness', '-')} "
                f"| {t.get('validity', '-')} "
                f"| {t.get('consistency', '-')} "
                f"| {t.get('coverage_estimate', '-')} |"
            )
        lines.append("")

        # Detalle de reglas existentes
        lines.append("## Detalle de Reglas Existentes")
        lines.append("")
        for t in tables:
            rules = t.get("rules", [])
            if not rules:
                continue
            lines.append(f"### {t['name']}")
            lines.append("")
            lines.append("| Regla | Dimension | Estado | % Pass | Descripcion |")
            lines.append("|-------|-----------|--------|--------|-------------|")
            for r in rules:
                pass_pct = f"{r['pass_pct']:.1f}%" if r.get("pass_pct") is not None else "-"
                status = r.get("status", "-")
                status_marker = "**" if status in ("KO", "WARNING") else ""
                lines.append(
                    f"| {status_marker}{r['name']}{status_marker} "
                    f"| {r.get('dimension', '-')} "
                    f"| {status_marker}{status}{status_marker} "
                    f"| {pass_pct} "
                    f"| {r.get('description', '-')} |"
                )
            lines.append("")

        # Gaps identificados
        all_gaps = []
        for t in tables:
            for g in t.get("gaps", []):
                all_gaps.append({**g, "table": t["name"]})

        if all_gaps:
            lines.append("## Gaps Identificados")
            lines.append("")
            # Orden: CRITICO > ALTO > MEDIO > BAJO
            priority_order = {"CRITICO": 0, "ALTO": 1, "MEDIO": 2, "BAJO": 3}
            all_gaps.sort(key=lambda x: priority_order.get(x.get("priority", "BAJO"), 99))
            lines.append("| Prioridad | Tabla | Columna | Dimension | Descripcion |")
            lines.append("|-----------|-------|---------|-----------|-------------|")
            for g in all_gaps:
                lines.append(
                    f"| {g.get('priority', '-')} "
                    f"| {g.get('table', '-')} "
                    f"| {g.get('column', '-')} "
                    f"| {g.get('dimension', '-')} "
                    f"| {g.get('description', '-')} |"
                )
            lines.append("")

    # Reglas creadas esta sesion
    rules_created = data.get("rules_created", [])
    if rules_created:
        lines.append("## Reglas Creadas en Esta Sesion")
        lines.append("")
        lines.append("| Regla | Tabla | Dimension | Estado |")
        lines.append("|-------|-------|-----------|--------|")
        for r in rules_created:
            lines.append(
                f"| {r['name']} "
                f"| {r.get('table', '-')} "
                f"| {r.get('dimension', '-')} "
                f"| {r.get('status', '-')} |"
            )
        lines.append("")

    # Recomendaciones
    recommendations = data.get("recommendations", [])
    if recommendations:
        lines.append("## Recomendaciones y Proximos Pasos")
        lines.append("")
        for i, rec in enumerate(recommendations, 1):
            lines.append(f"{i}. {rec}")
        lines.append("")

    lines.append("---")
    lines.append("*Generado por Data Quality Agent*")
    lines.append("")

    return "\n".join(lines)


def generate_markdown(data: dict, output_path: str) -> None:
    content = build_markdown(data)
    ensure_output_dir(output_path)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"[OK] Markdown generado: {output_path}")


# ---------------------------------------------------------------------------
# Generacion de PDF via WeasyPrint
# ---------------------------------------------------------------------------

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<style>
  body { font-family: Arial, sans-serif; font-size: 11pt; color: #1a1a2e; margin: 40px; }
  h1 { color: #16213e; border-bottom: 3px solid #0f3460; padding-bottom: 8px; font-size: 20pt; }
  h2 { color: #16213e; border-left: 4px solid #0f3460; padding-left: 10px; font-size: 14pt; margin-top: 30px; }
  h3 { color: #0f3460; font-size: 12pt; margin-top: 20px; }
  .meta { color: #555; font-size: 10pt; margin-bottom: 20px; }
  .meta span { margin-right: 20px; }
  table { width: 100%; border-collapse: collapse; margin: 10px 0 20px 0; font-size: 10pt; }
  th { background-color: #0f3460; color: white; padding: 7px 10px; text-align: left; }
  td { padding: 6px 10px; border-bottom: 1px solid #ddd; }
  tr:nth-child(even) { background-color: #f8f9fa; }
  .status-ok { color: #28a745; font-weight: bold; }
  .status-ko { color: #dc3545; font-weight: bold; }
  .status-warning { color: #ffc107; font-weight: bold; }
  .status-gap { color: #6c757d; }
  .priority-critico { color: #dc3545; font-weight: bold; }
  .priority-alto { color: #fd7e14; font-weight: bold; }
  .priority-medio { color: #ffc107; }
  .priority-bajo { color: #6c757d; }
  .summary-box { background: #f0f4ff; border: 1px solid #c0d0f0; border-radius: 6px;
                 padding: 15px 20px; margin: 20px 0; }
  .summary-box table { margin: 0; }
  .summary-box th { background-color: #0f3460; }
  .rec-list { padding-left: 20px; }
  .rec-list li { margin-bottom: 6px; }
  .footer { text-align: center; font-size: 9pt; color: #888; border-top: 1px solid #ddd;
            margin-top: 40px; padding-top: 10px; }
  @page { margin: 2cm; }
</style>
</head>
<body>
{{ content_html }}
<div class="footer">Generado por Data Quality Agent | {{ generated_at }}</div>
</body>
</html>
"""


def _status_class(status: str) -> str:
    mapping = {"OK": "status-ok", "KO": "status-ko", "WARNING": "status-warning",
               "Gap": "status-gap", "Parcial": "status-gap"}
    return mapping.get(status, "")


def _priority_class(priority: str) -> str:
    mapping = {"CRITICO": "priority-critico", "ALTO": "priority-alto",
               "MEDIO": "priority-medio", "BAJO": "priority-bajo"}
    return mapping.get(priority, "")


def build_html_body(data: dict) -> str:
    s = data.get("summary", {})
    parts = []

    # Cabecera
    parts.append(f"<h1>{data.get('title', 'Informe de Cobertura de Calidad del Dato')}</h1>")
    parts.append(
        f"<div class='meta'>"
        f"<span><b>Dominio:</b> {data.get('domain', '-')}</span>"
        f"<span><b>Scope:</b> {data.get('scope', '-')}</span>"
        f"<span><b>Fecha:</b> {data.get('generated_at', str(date.today()))}</span>"
        f"</div>"
    )

    # Resumen ejecutivo
    parts.append("<h2>Resumen Ejecutivo</h2>")
    parts.append("<div class='summary-box'>")
    parts.append("<table><tr><th>Metrica</th><th>Valor</th></tr>")
    summary_rows = [
        ("Tablas analizadas", s.get("tables_analyzed", "-")),
        ("Reglas existentes", s.get("rules_total", "-")),
        ("Estado OK / KO / WARNING", f"{s.get('rules_ok',0)} / {s.get('rules_ko',0)} / {s.get('rules_warning',0)}"),
        ("Sin ejecutar", s.get("rules_not_executed", "-")),
        ("Cobertura estimada", f"<b>{s.get('coverage_estimate', '-')}</b>"),
        ("Gaps criticos / moderados / bajos",
         f"{s.get('gaps_critical',0)} / {s.get('gaps_moderate',0)} / {s.get('gaps_low',0)}"),
    ]
    if s.get("rules_created_this_session"):
        summary_rows.append(("Reglas creadas esta sesion", s["rules_created_this_session"]))
    for label, value in summary_rows:
        parts.append(f"<tr><td>{label}</td><td>{value}</td></tr>")
    parts.append("</table></div>")

    # Cobertura por tabla
    tables = data.get("tables", [])
    if tables:
        parts.append("<h2>Cobertura por Tabla</h2>")
        parts.append("<table><tr><th>Tabla</th><th>Completeness</th><th>Uniqueness</th>"
                     "<th>Validity</th><th>Consistency</th><th>Cobertura</th></tr>")
        for t in tables:
            def cell(val):
                cls = _status_class(val or "-")
                return f"<td class='{cls}'>{val or '-'}</td>"
            parts.append(
                f"<tr><td><b>{t['name']}</b></td>"
                f"{cell(t.get('completeness'))}"
                f"{cell(t.get('uniqueness'))}"
                f"{cell(t.get('validity'))}"
                f"{cell(t.get('consistency'))}"
                f"<td><b>{t.get('coverage_estimate', '-')}</b></td></tr>"
            )
        parts.append("</table>")

        # Detalle de reglas
        parts.append("<h2>Detalle de Reglas Existentes</h2>")
        for t in tables:
            rules = t.get("rules", [])
            if not rules:
                continue
            parts.append(f"<h3>{t['name']}</h3>")
            parts.append("<table><tr><th>Regla</th><th>Dimension</th><th>Estado</th>"
                         "<th>% Pass</th><th>Descripcion</th></tr>")
            for r in rules:
                status = r.get("status", "-")
                cls = _status_class(status)
                pass_pct = f"{r['pass_pct']:.1f}%" if r.get("pass_pct") is not None else "-"
                parts.append(
                    f"<tr><td>{r['name']}</td><td>{r.get('dimension','-')}</td>"
                    f"<td class='{cls}'>{status}</td><td>{pass_pct}</td>"
                    f"<td>{r.get('description','-')}</td></tr>"
                )
            parts.append("</table>")

        # Gaps
        all_gaps = []
        for t in tables:
            for g in t.get("gaps", []):
                all_gaps.append({**g, "table": t["name"]})

        if all_gaps:
            priority_order = {"CRITICO": 0, "ALTO": 1, "MEDIO": 2, "BAJO": 3}
            all_gaps.sort(key=lambda x: priority_order.get(x.get("priority", "BAJO"), 99))
            parts.append("<h2>Gaps Identificados</h2>")
            parts.append("<table><tr><th>Prioridad</th><th>Tabla</th><th>Columna</th>"
                         "<th>Dimension</th><th>Descripcion</th></tr>")
            for g in all_gaps:
                pri = g.get("priority", "-")
                cls = _priority_class(pri)
                parts.append(
                    f"<tr><td class='{cls}'>{pri}</td><td>{g.get('table','-')}</td>"
                    f"<td><b>{g.get('column','-')}</b></td><td>{g.get('dimension','-')}</td>"
                    f"<td>{g.get('description','-')}</td></tr>"
                )
            parts.append("</table>")

    # Reglas creadas
    rules_created = data.get("rules_created", [])
    if rules_created:
        parts.append("<h2>Reglas Creadas en Esta Sesion</h2>")
        parts.append("<table><tr><th>Regla</th><th>Tabla</th><th>Dimension</th><th>Estado</th></tr>")
        for r in rules_created:
            status = r.get("status", "-")
            cls = "status-ok" if status == "created" else "status-ko"
            parts.append(
                f"<tr><td>{r['name']}</td><td>{r.get('table','-')}</td>"
                f"<td>{r.get('dimension','-')}</td><td class='{cls}'>{status}</td></tr>"
            )
        parts.append("</table>")

    # Recomendaciones
    recommendations = data.get("recommendations", [])
    if recommendations:
        parts.append("<h2>Recomendaciones y Proximos Pasos</h2>")
        parts.append("<ol class='rec-list'>")
        for rec in recommendations:
            parts.append(f"<li>{rec}</li>")
        parts.append("</ol>")

    return "\n".join(parts)


def generate_pdf(data: dict, output_path: str) -> None:
    try:
        from weasyprint import HTML as WeasyprintHTML
    except ImportError:
        print("[ERROR] weasyprint no esta instalado. Ejecuta: pip install weasyprint", file=sys.stderr)
        sys.exit(1)

    html_body = build_html_body(data)
    generated_at = data.get("generated_at", str(date.today()))
    html_full = HTML_TEMPLATE.replace("{{ content_html }}", html_body).replace(
        "{{ generated_at }}", generated_at
    )

    ensure_output_dir(output_path)
    WeasyprintHTML(string=html_full).write_pdf(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f"[OK] PDF generado: {output_path} ({size_kb} KB)")


# ---------------------------------------------------------------------------
# Generacion de DOCX via python-docx
# ---------------------------------------------------------------------------

def generate_docx(data: dict, output_path: str) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[ERROR] python-docx no esta instalado. Ejecuta: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    # Estilos basicos
    BLUE = RGBColor(15, 52, 96)    # #0f3460
    RED = RGBColor(220, 53, 69)
    GREEN = RGBColor(40, 167, 69)
    ORANGE = RGBColor(253, 126, 20)
    GRAY = RGBColor(108, 117, 125)

    s = data.get("summary", {})

    # Titulo
    title = doc.add_heading(data.get("title", "Informe de Cobertura de Calidad del Dato"), level=0)
    title.runs[0].font.color.rgb = BLUE

    # Meta
    meta = doc.add_paragraph()
    meta.add_run(f"Dominio: ").bold = True
    meta.add_run(f"{data.get('domain', '-')}   ")
    meta.add_run(f"Scope: ").bold = True
    meta.add_run(f"{data.get('scope', '-')}   ")
    meta.add_run(f"Fecha: ").bold = True
    meta.add_run(data.get("generated_at", str(date.today())))

    doc.add_heading("Resumen Ejecutivo", level=1)

    # Tabla de resumen
    summary_rows = [
        ("Tablas analizadas", str(s.get("tables_analyzed", "-"))),
        ("Reglas existentes", str(s.get("rules_total", "-"))),
        ("Estado OK / KO / WARNING",
         f"{s.get('rules_ok',0)} / {s.get('rules_ko',0)} / {s.get('rules_warning',0)}"),
        ("Sin ejecutar", str(s.get("rules_not_executed", "-"))),
        ("Cobertura estimada", str(s.get("coverage_estimate", "-"))),
        ("Gaps criticos / moderados / bajos",
         f"{s.get('gaps_critical',0)} / {s.get('gaps_moderate',0)} / {s.get('gaps_low',0)}"),
    ]
    if s.get("rules_created_this_session"):
        summary_rows.append(("Reglas creadas esta sesion", str(s["rules_created_this_session"])))

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metrica"
    hdr[1].text = "Valor"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    for label, value in summary_rows:
        row = tbl.add_row().cells
        row[0].text = label
        row[1].text = value

    # Cobertura por tabla
    tables = data.get("tables", [])
    if tables:
        doc.add_heading("Cobertura por Tabla", level=1)
        tbl2 = doc.add_table(rows=1, cols=6)
        tbl2.style = "Table Grid"
        headers = ["Tabla", "Completeness", "Uniqueness", "Validity", "Consistency", "Cobertura"]
        for i, h in enumerate(headers):
            tbl2.rows[0].cells[i].text = h
            tbl2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for t in tables:
            row = tbl2.add_row().cells
            vals = [t["name"], t.get("completeness", "-"), t.get("uniqueness", "-"),
                    t.get("validity", "-"), t.get("consistency", "-"), t.get("coverage_estimate", "-")]
            for i, v in enumerate(vals):
                row[i].text = v

        # Gaps
        all_gaps = []
        for t in tables:
            for g in t.get("gaps", []):
                all_gaps.append({**g, "table": t["name"]})
        if all_gaps:
            priority_order = {"CRITICO": 0, "ALTO": 1, "MEDIO": 2, "BAJO": 3}
            all_gaps.sort(key=lambda x: priority_order.get(x.get("priority", "BAJO"), 99))
            doc.add_heading("Gaps Identificados", level=1)
            tbl3 = doc.add_table(rows=1, cols=5)
            tbl3.style = "Table Grid"
            for i, h in enumerate(["Prioridad", "Tabla", "Columna", "Dimension", "Descripcion"]):
                tbl3.rows[0].cells[i].text = h
                tbl3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for g in all_gaps:
                row = tbl3.add_row().cells
                vals = [g.get("priority", "-"), g.get("table", "-"),
                        g.get("column", "-"), g.get("dimension", "-"), g.get("description", "-")]
                for i, v in enumerate(vals):
                    row[i].text = v

    # Reglas creadas
    rules_created = data.get("rules_created", [])
    if rules_created:
        doc.add_heading("Reglas Creadas en Esta Sesion", level=1)
        tbl4 = doc.add_table(rows=1, cols=4)
        tbl4.style = "Table Grid"
        for i, h in enumerate(["Regla", "Tabla", "Dimension", "Estado"]):
            tbl4.rows[0].cells[i].text = h
            tbl4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for r in rules_created:
            row = tbl4.add_row().cells
            for i, v in enumerate([r["name"], r.get("table", "-"), r.get("dimension", "-"), r.get("status", "-")]):
                row[i].text = v

    # Recomendaciones
    recommendations = data.get("recommendations", [])
    if recommendations:
        doc.add_heading("Recomendaciones y Proximos Pasos", level=1)
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")

    # Footer
    doc.add_paragraph("")
    footer_p = doc.add_paragraph("Generado por Data Quality Agent")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = GRAY

    ensure_output_dir(output_path)
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f"[OK] DOCX generado: {output_path} ({size_kb} KB)")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Genera informes de calidad del dato")
    parser.add_argument("--format", choices=["pdf", "docx", "md"], required=True,
                        help="Formato de salida: pdf, docx o md")
    parser.add_argument("--output", required=True, help="Ruta del archivo de salida")
    parser.add_argument("--input-json", dest="input_json", default=None,
                        help="JSON con los datos del informe (string)")
    parser.add_argument("--input-file", dest="input_file", default=None,
                        help="Ruta a un archivo JSON con los datos del informe")
    args = parser.parse_args()

    try:
        data = load_input(args.input_json, args.input_file)
    except (json.JSONDecodeError, ValueError) as e:
        print(f"[ERROR] No se pudieron cargar los datos de entrada: {e}", file=sys.stderr)
        sys.exit(1)

    if args.format == "md":
        generate_markdown(data, args.output)
    elif args.format == "pdf":
        generate_pdf(data, args.output)
    elif args.format == "docx":
        generate_docx(data, args.output)


if __name__ == "__main__":
    main()

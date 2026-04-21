"""DOCX Generator for BI/BA Analytics Agent.

Provides two rendering modes (parallel API to PDFGenerator):
- **scaffold**: Predefined structure (executive summary -> methodology ->
  data -> analysis -> conclusions) with KPIs, figures, tables, callouts.
- **render_from_markdown**: Converts markdown content to DOCX elements.

Usage (from a script that has added skills/analyze/report/tools to sys.path):
    from docx_generator import DOCXGenerator

    gen = DOCXGenerator(style="corporate")

    # Scaffold mode
    gen.render_scaffold(
        title="Sales Analysis Q4",
        kpis=[{"value": "1.2M", "label": "Revenue", "change": 15}],
        executive_summary="<p>Key findings...</p>",
        analysis="<p>Detailed analysis...</p>",
        conclusions="<p>Recommendations...</p>",
    )
    gen.save("output/report.docx")

    # Markdown mode
    gen.render_from_markdown(md_content, title="Report")
    gen.save("output/report.docx")
"""

import base64
import re
import tempfile
from datetime import date
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag

import markdown as md_lib
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from css_builder import get_palette, aesthetic_to_override_tokens
from i18n import get_labels

# A4 margins in cm
_MARGIN_CM = 2.5
_IMAGE_WIDTH_INCHES = 6.0

# Callout background lightening factor
_CALLOUT_LIGHTEN = 0.7


def _rgb_to_hex(rgb_tuple: tuple) -> str:
    """Convert (R, G, B) tuple to hex string without '#'."""
    return f"{rgb_tuple[0]:02X}{rgb_tuple[1]:02X}{rgb_tuple[2]:02X}"


def _lighten_color(rgb_tuple: tuple, factor: float = _CALLOUT_LIGHTEN) -> tuple:
    """Lighten an RGB color by blending with white."""
    return tuple(int(c + (255 - c) * factor) for c in rgb_tuple)


def _strip_html(html: str) -> str:
    """Extract plain text from HTML, stripping all tags."""
    return re.sub(r"<[^>]+>", "", html).strip()


class DOCXGenerator:
    """Generate professional DOCX reports with templates or markdown."""

    def __init__(self, style: str = "corporate", author: str | None = None,
                 aesthetic_direction: dict | None = None):
        override_tokens = aesthetic_to_override_tokens(aesthetic_direction)
        self._palette = get_palette(style, override_tokens=override_tokens or None)
        self._style_name = style
        self._author = author
        self._aesthetic = aesthetic_direction or {}
        self._doc: Document | None = None

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def render_scaffold(
        self,
        title: str,
        executive_summary: str = "",
        methodology: str = "",
        data_section: str = "",
        analysis: str = "",
        conclusions: str = "",
        kpis: list[dict] | None = None,
        figures: list[dict] | None = None,
        data_tables: list[dict] | None = None,
        callouts: list[dict] | None = None,
        domain: str | None = None,
        subtitle: str | None = None,
        author: str | None = None,
        show_cover: bool = True,
        lang: str | None = None,
        labels: dict[str, str] | None = None,
    ) -> "DOCXGenerator":
        """Render using scaffold structure (same signature as PDFGenerator).

        `lang` picks a language from the i18n catalogue (fallbacks to
        `.agent_lang` file, then English). `labels` is an override dict whose
        keys take precedence over the catalogue.
        """
        resolved_labels = get_labels(lang=lang, overrides=labels)

        doc = self._setup_document()
        self._apply_styles(doc)

        if show_cover:
            self._add_cover_page(doc, title, subtitle=subtitle,
                                 author=author or self._author,
                                 domain=domain,
                                 labels=resolved_labels)

        # KPIs
        if kpis:
            self._add_kpi_table(doc, kpis)

        # Sections — localised headings
        sections = [
            (resolved_labels["report.executive_summary"], executive_summary),
            (resolved_labels["report.methodology"], methodology),
            (resolved_labels["report.data_sources"], data_section),
            (resolved_labels["report.analysis"], analysis),
            (resolved_labels["report.conclusions"], conclusions),
        ]
        first_section_rendered = False
        for heading, content in sections:
            if content and content.strip():
                self._add_section_text(
                    doc, content, heading,
                    page_break_before=first_section_rendered,
                )
                first_section_rendered = True

        # Figures
        if figures:
            for fig in figures:
                self._add_figure(doc, fig)

        # Data tables
        if data_tables:
            for tbl in data_tables:
                self._add_data_table(doc, tbl)

        # Callouts
        if callouts:
            for callout in callouts:
                self._add_callout(doc, callout)

        self._doc = doc
        return self

    def render_from_markdown(
        self,
        md_content: str,
        title: str | None = None,
        domain: str | None = None,
        subtitle: str | None = None,
        author: str | None = None,
        show_cover: bool = True,
        lang: str | None = None,
        labels: dict[str, str] | None = None,
    ) -> "DOCXGenerator":
        """Render from markdown content.

        If `title` is not provided, falls back to the localised default
        ("Report" / "Informe" / ...). `lang` and `labels` behave like in
        `render_scaffold`.
        """
        resolved_labels = get_labels(lang=lang, overrides=labels)
        effective_title = title or resolved_labels["report.default_title"]

        doc = self._setup_document()
        self._apply_styles(doc)

        if show_cover:
            self._add_cover_page(doc, effective_title, subtitle=subtitle,
                                 author=author or self._author,
                                 domain=domain,
                                 labels=resolved_labels)

        elements = self._parse_markdown_to_elements(md_content)
        self._render_elements(doc, elements)

        self._doc = doc
        return self

    def save(self, output_path: str | Path) -> Path:
        """Save the document to disk."""
        if self._doc is None:
            raise RuntimeError(
                "Call render_scaffold() or render_from_markdown() first."
            )
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Document setup
    # ------------------------------------------------------------------

    def _setup_document(self) -> Document:
        """Create a Document with A4 margins."""
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        margin = Cm(_MARGIN_CM)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
        return doc

    def _apply_styles(self, doc: Document) -> None:
        """Configure document styles using palette colors and fonts."""
        p = self._palette
        font_main = p.get("font_main", "Arial")
        # When ``aesthetic_direction`` supplied a font_pair, the first
        # element is the display face — use it for headings so the document
        # honours the design-first decision.
        fp = self._aesthetic.get("font_pair")
        if (fp and isinstance(fp, list) and len(fp) == 2
                and isinstance(fp[0], str) and fp[0]):
            heading_font = fp[0]
        else:
            heading_font = font_main
        primary = p.get("primary", (26, 54, 93))
        primary_light = p.get("primary_light", (70, 130, 180))
        text_color = p.get("text", (33, 33, 33))
        text_muted = p.get("text_muted", (100, 100, 100))

        # Normal
        style_normal = doc.styles["Normal"]
        style_normal.font.name = font_main
        style_normal.font.size = Pt(11)
        style_normal.font.color.rgb = RGBColor(*text_color)
        style_normal.paragraph_format.line_spacing = 1.15
        style_normal.paragraph_format.widow_control = True

        # Headings
        for level, (size, color) in {
            1: (24, primary),
            2: (18, primary),
            3: (14, primary_light),
        }.items():
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.name = heading_font
                style.font.size = Pt(size)
                style.font.bold = True
                style.font.color.rgb = RGBColor(*color)
                style.paragraph_format.keep_with_next = True

        # Caption (use Subtitle style as proxy)
        if "Caption" in doc.styles:
            cap_style = doc.styles["Caption"]
            cap_style.font.name = font_main
            cap_style.font.size = Pt(9)
            cap_style.font.italic = True
            cap_style.font.color.rgb = RGBColor(*text_muted)

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _add_cover_page(self, doc: Document, title: str,
                        subtitle: str | None = None,
                        author: str | None = None,
                        domain: str | None = None,
                        labels: dict[str, str] | None = None) -> None:
        """Add a cover page with title, metadata, and page break."""
        if labels is None:
            labels = get_labels()
        p = self._palette
        primary = p.get("primary", (26, 54, 93))

        # Spacer
        for _ in range(6):
            doc.add_paragraph("")

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*primary)
        run.font.name = p.get("font_main", "Arial")

        # Subtitle
        if subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = Pt(16)
            sub_run.font.color.rgb = RGBColor(*p.get("text_muted", (100, 100, 100)))

        doc.add_paragraph("")

        # Metadata — localised labels
        meta_lines = []
        if author:
            meta_lines.append(f"{labels['cover.author']}: {author}")
        if domain:
            meta_lines.append(f"{labels['cover.domain']}: {domain}")
        meta_lines.append(f"{labels['cover.date']}: {date.today().strftime('%d/%m/%Y')}")

        for line in meta_lines:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_run = meta_para.add_run(line)
            meta_run.font.size = Pt(12)
            meta_run.font.color.rgb = RGBColor(*p.get("text_muted", (100, 100, 100)))

        # Page break
        doc.add_page_break()

    # ------------------------------------------------------------------
    # KPIs
    # ------------------------------------------------------------------

    def _add_kpi_table(self, doc: Document, kpis: list[dict]) -> None:
        """Add a horizontal KPI table: values row, labels row, change row."""
        p = self._palette
        primary = p.get("primary", (26, 54, 93))
        success = p.get("success", (56, 161, 105))
        danger = p.get("danger", (229, 62, 62))
        text_muted = p.get("text_muted", (100, 100, 100))

        n = len(kpis)
        if n == 0:
            return

        table = doc.add_table(rows=3, cols=n)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Remove borders
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell, top=None, bottom=None, left=None, right=None)

        # Row 0: Values
        for i, kpi in enumerate(kpis):
            cell = table.rows[0].cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(kpi.get("value", "")))
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*primary)

        # Row 1: Labels
        for i, kpi in enumerate(kpis):
            cell = table.rows[1].cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(kpi.get("label", "")))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*text_muted)

        # Row 2: Change indicators
        for i, kpi in enumerate(kpis):
            change = kpi.get("change")
            cell = table.rows[2].cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if change is not None:
                sign = "+" if change > 0 else ""
                run = para.add_run(f"{sign}{change}%")
                run.font.size = Pt(10)
                run.font.bold = True
                color = success if change >= 0 else danger
                run.font.color.rgb = RGBColor(*color)

        # Pagination: prevent KPI rows from splitting across pages
        for row in table.rows:
            self._set_row_cant_split(row)

        doc.add_paragraph("")  # spacing

    # ------------------------------------------------------------------
    # Figures
    # ------------------------------------------------------------------

    def _add_figure(self, doc: Document, fig: dict) -> None:
        """Add an image with caption. Fallback if image not found."""
        path = fig.get("path", "")
        caption = fig.get("caption", "")

        if path and Path(path).is_file():
            doc.add_picture(str(path), width=Inches(_IMAGE_WIDTH_INCHES))
            # Center the last paragraph (the image) and keep with caption
            img_para = doc.paragraphs[-1]
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.keep_with_next = True
            img_para.paragraph_format.keep_together = True
        else:
            p = doc.add_paragraph(f"[Image not found: {path}]")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
            p.paragraph_format.keep_with_next = True

        if caption:
            cap_para = doc.add_paragraph(caption, style="Caption")
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")  # spacing

    # ------------------------------------------------------------------
    # Data tables
    # ------------------------------------------------------------------

    def _add_data_table(self, doc: Document, table_data: dict) -> None:
        """Add a data table with colored header and alternating rows."""
        p = self._palette
        primary = p.get("primary", (26, 54, 93))
        bg_alt = p.get("bg_alt", (245, 245, 245))

        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        caption = table_data.get("caption", "")

        if not headers:
            return

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            self._shade_cell(cell, _rgb_to_hex(primary))
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(header))
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)

        # Pagination: repeat header row on each page, prevent row splitting
        self._set_repeat_header_row(table.rows[0])
        self._set_row_cant_split(table.rows[0])

        # Data rows
        for row_idx, row_data in enumerate(rows):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                if row_idx % 2 == 1:
                    self._shade_cell(cell, _rgb_to_hex(bg_alt))
                para = cell.paragraphs[0]
                run = para.add_run(str(value))
                run.font.size = Pt(10)
            self._set_row_cant_split(table.rows[row_idx + 1])

        if caption:
            cap_para = doc.add_paragraph(caption, style="Caption")
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")  # spacing

    # ------------------------------------------------------------------
    # Callouts
    # ------------------------------------------------------------------

    def _add_callout(self, doc: Document, callout: dict) -> None:
        """Add a callout as a 1-cell table with background shading."""
        p = self._palette
        callout_type = callout.get("type", "info")
        content = callout.get("content", "")

        color_map = {
            "info": p.get("accent", (70, 130, 180)),
            "success": p.get("success", (56, 161, 105)),
            "warning": p.get("warning", (214, 158, 46)),
            "danger": p.get("danger", (229, 62, 62)),
        }
        base_color = color_map.get(callout_type, color_map["info"])
        bg_color = _lighten_color(base_color)

        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        self._shade_cell(cell, _rgb_to_hex(bg_color))

        # Set cell width to approximate page width
        cell.width = Inches(_IMAGE_WIDTH_INCHES)

        para = cell.paragraphs[0]
        para.paragraph_format.keep_together = True
        # Parse basic HTML if present
        plain = _strip_html(content) if "<" in content else content
        run = para.add_run(plain)
        run.font.size = Pt(10)

        # Pagination: prevent callout row from splitting across pages
        self._set_row_cant_split(table.rows[0])

        doc.add_paragraph("")  # spacing

    # ------------------------------------------------------------------
    # Section text (HTML-aware)
    # ------------------------------------------------------------------

    def _add_section_text(self, doc: Document, text: str, heading: str,
                          page_break_before: bool = False) -> None:
        """Add a heading + content. Parses basic HTML tags."""
        h = doc.add_heading(heading, level=2)
        if page_break_before:
            h.paragraph_format.page_break_before = True
        self._render_html_content(doc, text)

    @staticmethod
    def _base64_to_temp_file(data_uri: str) -> str:
        """Decode a base64 data URI to a temporary PNG file and return its path."""
        # Strip the data:image/...;base64, prefix
        if "," in data_uri:
            data_uri = data_uri.split(",", 1)[1]
        img_bytes = base64.b64decode(data_uri)
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.write(img_bytes)
        tmp.close()
        return tmp.name

    def _render_html_content(self, doc: Document, html_text: str) -> None:
        """Parse HTML with BeautifulSoup and render all elements inline."""
        soup = BeautifulSoup(html_text, "html.parser")
        for element in soup.children:
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text:
                    para = doc.add_paragraph()
                    para.add_run(text)
                continue

            if not isinstance(element, Tag):
                continue

            tag = element.name

            if tag in ("h3", "h4"):
                level = int(tag[1])
                h = doc.add_heading(element.get_text(strip=True), level=level)
                h.paragraph_format.keep_with_next = True

            elif tag == "p":
                para = doc.add_paragraph()
                self._add_formatted_runs(para, str(element.decode_contents()))

            elif tag == "figure":
                img_tag = element.find("img")
                figcaption = element.find("figcaption")
                caption = figcaption.get_text(strip=True) if figcaption else ""
                if img_tag:
                    src = img_tag.get("src", "")
                    alt = img_tag.get("alt", "")
                    caption = caption or alt
                    if src.startswith("data:"):
                        tmp_path = self._base64_to_temp_file(src)
                        try:
                            self._add_figure(doc, {"path": tmp_path, "caption": caption})
                        finally:
                            Path(tmp_path).unlink(missing_ok=True)
                    else:
                        self._add_figure(doc, {"path": src, "caption": caption})

            elif tag == "img":
                src = element.get("src", "")
                alt = element.get("alt", "")
                if src.startswith("data:"):
                    tmp_path = self._base64_to_temp_file(src)
                    try:
                        self._add_figure(doc, {"path": tmp_path, "caption": alt})
                    finally:
                        Path(tmp_path).unlink(missing_ok=True)
                else:
                    self._add_figure(doc, {"path": src, "caption": alt})

            elif tag == "table":
                self._render_html_table(doc, str(element))

            elif tag in ("ul", "ol"):
                items_tags = element.find_all("li")
                items = [li.get_text(strip=True) for li in items_tags]
                self._add_list(doc, items, ordered=(tag == "ol"))

            elif tag == "div" and "callout" in element.get("class", []):
                callout_type = "info"
                for cls in element.get("class", []):
                    if cls in ("success", "warning", "danger", "info"):
                        callout_type = cls
                self._add_callout(doc, {
                    "type": callout_type,
                    "content": element.get_text(strip=True),
                })

            elif tag == "blockquote":
                self._add_callout(doc, {
                    "type": "info",
                    "content": element.get_text(strip=True),
                })

            elif tag == "br":
                doc.add_paragraph("")

            else:
                # Fallback: extract text and render as paragraph
                text = element.get_text(strip=True)
                if text:
                    para = doc.add_paragraph()
                    self._add_formatted_runs(para, str(element.decode_contents()))

    def _add_formatted_runs(self, para, html: str) -> None:
        """Add runs with bold/italic from HTML inline tags."""
        # Pattern to match <strong>, <b>, <em>, <i> tags
        pattern = r"(<strong>|<b>)(.*?)(</strong>|</b>)|(<em>|<i>)(.*?)(</em>|</i>)"
        pos = 0
        for m in re.finditer(pattern, html):
            # Text before this match
            before = _strip_html(html[pos:m.start()])
            if before:
                para.add_run(before)

            if m.group(2) is not None:
                # Bold
                run = para.add_run(_strip_html(m.group(2)))
                run.font.bold = True
            elif m.group(5) is not None:
                # Italic
                run = para.add_run(_strip_html(m.group(5)))
                run.font.italic = True

            pos = m.end()

        # Remaining text
        remaining = _strip_html(html[pos:])
        if remaining:
            para.add_run(remaining)

    # ------------------------------------------------------------------
    # Markdown parser
    # ------------------------------------------------------------------

    def _parse_markdown_to_elements(self, md_content: str) -> list[dict]:
        """Convert markdown to an intermediate list of elements.

        Converts md -> HTML (via markdown lib) -> parses HTML tags into
        element dicts that _render_elements() can process.
        """
        html = md_lib.markdown(md_content, extensions=[
            "extra", "codehilite", "toc", "sane_lists",
        ])
        return self._html_to_elements(html)

    def _html_to_elements(self, html: str) -> list[dict]:
        """Parse HTML into a list of element dicts using BeautifulSoup."""
        soup = BeautifulSoup(html, "html.parser")
        elements = []

        for tag in soup.children:
            if isinstance(tag, NavigableString):
                text = str(tag).strip()
                if text:
                    elements.append({"type": "paragraph", "html": text})
                continue

            if not isinstance(tag, Tag):
                continue

            name = tag.name

            if name in ("h1", "h2", "h3"):
                elements.append({
                    "type": "heading",
                    "level": int(name[1]),
                    "text": tag.get_text(strip=True),
                })

            elif name == "p":
                inner = str(tag.decode_contents()).strip()
                # Check if paragraph contains only an image
                img = tag.find("img")
                if img and tag.get_text(strip=True) == "":
                    src = img.get("src", "")
                    alt = img.get("alt", "")
                    elements.append({
                        "type": "image",
                        "path": src,
                        "caption": alt,
                    })
                elif inner:
                    elements.append({"type": "paragraph", "html": inner})

            elif name == "figure":
                img = tag.find("img")
                figcaption = tag.find("figcaption")
                caption = figcaption.get_text(strip=True) if figcaption else ""
                if img:
                    src = img.get("src", "")
                    alt = img.get("alt", "")
                    elements.append({
                        "type": "image",
                        "path": src,
                        "caption": caption or alt,
                    })

            elif name == "img":
                elements.append({
                    "type": "image",
                    "path": tag.get("src", ""),
                    "caption": tag.get("alt", ""),
                })

            elif name == "table":
                elements.append({
                    "type": "table",
                    "html": str(tag),
                })

            elif name == "pre":
                code_tag = tag.find("code")
                text = code_tag.get_text() if code_tag else tag.get_text()
                elements.append({"type": "code", "text": text})

            elif name == "blockquote":
                elements.append({
                    "type": "blockquote",
                    "text": tag.get_text(strip=True),
                })

            elif name in ("ul", "ol"):
                items = [li.get_text(strip=True) for li in tag.find_all("li")]
                elements.append({
                    "type": "list",
                    "ordered": name == "ol",
                    "items": items,
                })

            elif name == "div" and "callout" in tag.get("class", []):
                callout_type = "info"
                for cls in tag.get("class", []):
                    if cls in ("success", "warning", "danger", "info"):
                        callout_type = cls
                elements.append({
                    "type": "callout",
                    "callout_type": callout_type,
                    "content": tag.get_text(strip=True),
                })

        return elements

    def _render_elements(self, doc: Document, elements: list[dict]) -> None:
        """Render parsed elements into the Document."""
        for el in elements:
            t = el["type"]
            if t == "heading":
                h = doc.add_heading(el["text"], level=el["level"])
                h.paragraph_format.keep_with_next = True
            elif t == "paragraph":
                para = doc.add_paragraph()
                self._add_formatted_runs(para, el["html"])
            elif t == "image":
                path = el["path"]
                caption = el.get("caption", "")
                if path.startswith("data:"):
                    tmp_path = self._base64_to_temp_file(path)
                    try:
                        self._add_figure(doc, {"path": tmp_path, "caption": caption})
                    finally:
                        Path(tmp_path).unlink(missing_ok=True)
                else:
                    self._add_figure(doc, {"path": path, "caption": caption})
            elif t == "table":
                self._render_html_table(doc, el["html"])
            elif t == "code":
                self._add_code_block(doc, el["text"])
            elif t == "blockquote":
                self._add_callout(doc, {"type": "info", "content": el["text"]})
            elif t == "callout":
                self._add_callout(doc, {"type": el["callout_type"], "content": el["content"]})
            elif t == "list":
                self._add_list(doc, el["items"], ordered=el["ordered"])

    def _render_html_table(self, doc: Document, table_html: str) -> None:
        """Parse an HTML table and render it as a DOCX table."""
        # Extract headers from <th> tags
        headers = re.findall(r"<th[^>]*>(.*?)</th>", table_html, re.DOTALL)
        headers = [_strip_html(h) for h in headers]

        # Extract rows
        row_matches = re.findall(r"<tr>(.*?)</tr>", table_html, re.DOTALL)
        rows = []
        for row_html in row_matches:
            cells = re.findall(r"<td[^>]*>(.*?)</td>", row_html, re.DOTALL)
            if cells:
                rows.append([_strip_html(c) for c in cells])

        if headers:
            self._add_data_table(doc, {"headers": headers, "rows": rows})
        elif rows:
            # No headers, use first row as header
            self._add_data_table(doc, {"headers": rows[0], "rows": rows[1:]})

    def _add_code_block(self, doc: Document, text: str) -> None:
        """Add a code block with monospace font and gray background."""
        p = self._palette
        font_mono = p.get("font_mono", "Consolas")

        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        self._shade_cell(cell, "F0F0F0")
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.name = font_mono
        run.font.size = Pt(9)

        # Pagination: prevent code block row from splitting across pages
        self._set_row_cant_split(table.rows[0])

        doc.add_paragraph("")  # spacing

    def _add_list(self, doc: Document, items: list[str],
                  ordered: bool = False) -> None:
        """Add a bulleted or numbered list."""
        for item in items:
            style = "List Number" if ordered else "List Bullet"
            doc.add_paragraph(item, style=style)

    # ------------------------------------------------------------------
    # Cell shading helper
    # ------------------------------------------------------------------

    @staticmethod
    def _shade_cell(cell, hex_color: str) -> None:
        """Apply background shading to a table cell."""
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)

    @staticmethod
    def _set_row_cant_split(row) -> None:
        """Prevent a table row from splitting across pages."""
        trPr = row._tr.get_or_add_trPr()
        cant_split = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
        trPr.append(cant_split)

    @staticmethod
    def _set_repeat_header_row(row) -> None:
        """Repeat this row as header on every page when table spans pages."""
        trPr = row._tr.get_or_add_trPr()
        tbl_header = parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
        trPr.append(tbl_header)

    @staticmethod
    def _set_cell_border(cell, **kwargs) -> None:
        """Set or remove cell borders. Pass None to remove a border."""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = parse_xml(f"<w:tcBorders {nsdecls('w')}/>")
            tcPr.append(tcBorders)
        for edge in ("top", "bottom", "left", "right"):
            if edge in kwargs:
                element = tcBorders.find(qn(f"w:{edge}"))
                if element is not None:
                    tcBorders.remove(element)
                if kwargs[edge] is None:
                    # Remove border
                    new_el = parse_xml(
                        f'<w:{edge} {nsdecls("w")} w:val="none" w:sz="0" w:space="0"/>'
                    )
                    tcBorders.append(new_el)

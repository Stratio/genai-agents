"""Tests for DOCXGenerator — scaffold, markdown, styles, callouts, images."""

import sys
import tempfile
from pathlib import Path

import pytest
from docx import Document

# Ensure tools/ is on the path
sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx_generator import DOCXGenerator, _rgb_to_hex, _lighten_color, _strip_html


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def tmp_dir():
    """Provide a temporary directory for test outputs."""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def sample_kpis():
    return [
        {"value": "1.2M", "label": "Revenue", "change": 15},
        {"value": "845", "label": "Customers", "change": -3},
        {"value": "72%", "label": "Retention", "change": 0},
    ]


@pytest.fixture
def sample_figures(tmp_dir):
    """Create a minimal PNG file for testing."""
    # Create a tiny valid PNG (1x1 white pixel)
    import struct
    import zlib

    def _create_minimal_png(path):
        sig = b'\x89PNG\r\n\x1a\n'
        ihdr_data = struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)
        ihdr_crc = zlib.crc32(b'IHDR' + ihdr_data) & 0xffffffff
        ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data + struct.pack('>I', ihdr_crc)
        raw = zlib.compress(b'\x00\xff\xff\xff')
        idat_crc = zlib.crc32(b'IDAT' + raw) & 0xffffffff
        idat = struct.pack('>I', len(raw)) + b'IDAT' + raw + struct.pack('>I', idat_crc)
        iend_crc = zlib.crc32(b'IEND') & 0xffffffff
        iend = struct.pack('>I', 0) + b'IEND' + struct.pack('>I', iend_crc)
        with open(path, 'wb') as f:
            f.write(sig + ihdr + idat + iend)

    img_path = tmp_dir / "test_chart.png"
    _create_minimal_png(str(img_path))
    return [{"path": str(img_path), "caption": "Test Chart"}]


@pytest.fixture
def sample_tables():
    return [
        {
            "headers": ["Region", "Revenue", "Growth"],
            "rows": [
                ["North", "500K", "+12%"],
                ["South", "350K", "-5%"],
                ["East", "420K", "+8%"],
            ],
            "caption": "Regional performance",
        }
    ]


@pytest.fixture
def sample_callouts():
    return [
        {"type": "info", "content": "This is an informational note."},
        {"type": "success", "content": "Target exceeded by 15%."},
        {"type": "warning", "content": "Data quality issues detected."},
        {"type": "danger", "content": "Critical anomaly in Q4."},
    ]


# ---------------------------------------------------------------------------
# Helper tests
# ---------------------------------------------------------------------------

class TestHelpers:
    def test_rgb_to_hex(self):
        assert _rgb_to_hex((26, 54, 93)) == "1A365D"
        assert _rgb_to_hex((0, 0, 0)) == "000000"
        assert _rgb_to_hex((255, 255, 255)) == "FFFFFF"

    def test_lighten_color(self):
        result = _lighten_color((100, 100, 100), 0.5)
        assert all(100 < v < 256 for v in result)

    def test_strip_html(self):
        assert _strip_html("<p>Hello <strong>world</strong></p>") == "Hello world"
        assert _strip_html("plain text") == "plain text"


# ---------------------------------------------------------------------------
# Scaffold tests
# ---------------------------------------------------------------------------

class TestScaffold:
    def test_generates_valid_docx(self, tmp_dir, sample_kpis):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Test Report",
            executive_summary="<p>Key findings here.</p>",
            methodology="<p>We analyzed data.</p>",
            analysis="<p>Results show growth.</p>",
            conclusions="<p>Recommendations follow.</p>",
            kpis=sample_kpis,
        )
        path = gen.save(tmp_dir / "report.docx")
        assert path.exists()
        assert path.stat().st_size > 0

        doc = Document(str(path))
        assert len(doc.paragraphs) > 0

    def test_has_cover_page(self, tmp_dir):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Cover Test",
            author="Test Author",
            domain="TestDomain",
            show_cover=True,
        )
        path = gen.save(tmp_dir / "cover.docx")
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        assert any("Cover Test" in t for t in texts)
        assert any("Test Author" in t for t in texts)

    def test_no_cover_when_disabled(self, tmp_dir):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="No Cover",
            show_cover=False,
            executive_summary="<p>Content only.</p>",
        )
        path = gen.save(tmp_dir / "nocover.docx")
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        # Should not have the large title paragraph from cover
        assert not any("Autor:" in t for t in texts)

    def test_has_kpi_table(self, tmp_dir, sample_kpis):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(title="KPI Test", kpis=sample_kpis)
        path = gen.save(tmp_dir / "kpi.docx")
        doc = Document(str(path))
        assert len(doc.tables) >= 1
        # Check KPI values are in the table
        table = doc.tables[0]
        cell_texts = [cell.text for row in table.rows for cell in row.cells]
        assert "1.2M" in cell_texts
        assert "845" in cell_texts

    def test_has_all_sections(self, tmp_dir):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Sections Test",
            executive_summary="<p>Summary</p>",
            methodology="<p>Method</p>",
            data_section="<p>Data</p>",
            analysis="<p>Analysis</p>",
            conclusions="<p>Conclusions</p>",
            show_cover=False,
        )
        path = gen.save(tmp_dir / "sections.docx")
        doc = Document(str(path))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Resumen Ejecutivo" in headings
        assert "Metodologia" in headings
        assert "Datos y Fuentes" in headings
        assert "Analisis" in headings
        assert "Conclusiones" in headings

    def test_has_figures(self, tmp_dir, sample_figures):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Figure Test",
            figures=sample_figures,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "figures.docx")
        doc = Document(str(path))
        # Check that "Test Chart" caption is present
        texts = [p.text for p in doc.paragraphs]
        assert any("Test Chart" in t for t in texts)

    def test_has_data_tables(self, tmp_dir, sample_tables):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Table Test",
            data_tables=sample_tables,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "tables.docx")
        doc = Document(str(path))
        # Should have at least one table with headers
        found = False
        for table in doc.tables:
            cell_texts = [cell.text for row in table.rows for cell in row.cells]
            if "Region" in cell_texts and "Revenue" in cell_texts:
                found = True
                break
        assert found


# ---------------------------------------------------------------------------
# Markdown (free-form) tests
# ---------------------------------------------------------------------------

class TestMarkdown:
    def test_parses_headings(self, tmp_dir):
        md = "# Main Title\n\n## Section One\n\nSome text here.\n\n### Subsection\n\nMore text."
        gen = DOCXGenerator(style="corporate")
        gen.render_from_markdown(md, title="MD Test", show_cover=False)
        path = gen.save(tmp_dir / "md.docx")
        doc = Document(str(path))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Main Title" in headings
        assert "Section One" in headings
        assert "Subsection" in headings

    def test_parses_paragraphs(self, tmp_dir):
        md = "First paragraph.\n\nSecond paragraph with **bold** text."
        gen = DOCXGenerator(style="corporate")
        gen.render_from_markdown(md, title="Para Test", show_cover=False)
        path = gen.save(tmp_dir / "para.docx")
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        assert any("First paragraph" in t for t in texts)

    def test_parses_tables(self, tmp_dir):
        md = "| Col A | Col B |\n|-------|-------|\n| 1 | 2 |\n| 3 | 4 |"
        gen = DOCXGenerator(style="corporate")
        gen.render_from_markdown(md, title="Table Test", show_cover=False)
        path = gen.save(tmp_dir / "mdtable.docx")
        doc = Document(str(path))
        assert len(doc.tables) >= 1

    def test_parses_lists(self, tmp_dir):
        md = "- Item one\n- Item two\n- Item three"
        gen = DOCXGenerator(style="corporate")
        gen.render_from_markdown(md, title="List Test", show_cover=False)
        path = gen.save(tmp_dir / "mdlist.docx")
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        assert any("Item one" in t for t in texts)


# ---------------------------------------------------------------------------
# Style tests
# ---------------------------------------------------------------------------

class TestStyles:
    @pytest.mark.parametrize("style", ["corporate", "modern", "academic"])
    def test_all_styles_generate(self, tmp_dir, style):
        gen = DOCXGenerator(style=style)
        gen.render_scaffold(title=f"Style: {style}", show_cover=False)
        path = gen.save(tmp_dir / f"{style}.docx")
        assert path.exists()
        assert path.stat().st_size > 0

    def test_unknown_style_falls_back_to_corporate(self, tmp_dir):
        gen = DOCXGenerator(style="nonexistent")
        gen.render_scaffold(title="Fallback Test", show_cover=False)
        path = gen.save(tmp_dir / "fallback.docx")
        assert path.exists()


# ---------------------------------------------------------------------------
# Callout tests
# ---------------------------------------------------------------------------

class TestCallouts:
    @pytest.mark.parametrize("callout_type", ["info", "success", "warning", "danger"])
    def test_callout_generates_table(self, tmp_dir, callout_type):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Callout Test",
            callouts=[{"type": callout_type, "content": f"Test {callout_type}"}],
            show_cover=False,
        )
        path = gen.save(tmp_dir / f"callout_{callout_type}.docx")
        doc = Document(str(path))
        # Callout is rendered as a 1-cell table
        assert len(doc.tables) >= 1
        cell_texts = [cell.text for row in doc.tables[-1].rows for cell in row.cells]
        assert any(callout_type in t.lower() or f"Test {callout_type}" in t for t in cell_texts)


# ---------------------------------------------------------------------------
# Image tests
# ---------------------------------------------------------------------------

class TestImages:
    def test_missing_image_no_exception(self, tmp_dir):
        """Image not found should degrade gracefully."""
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Missing Image",
            figures=[{"path": "/nonexistent/path.png", "caption": "Ghost"}],
            show_cover=False,
        )
        path = gen.save(tmp_dir / "missing.docx")
        assert path.exists()
        doc = Document(str(path))
        texts = [p.text for p in doc.paragraphs]
        assert any("Imagen no encontrada" in t for t in texts)


# ---------------------------------------------------------------------------
# save() error handling
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Inline rendering tests
# ---------------------------------------------------------------------------

class TestInlineRendering:
    """Test that images, tables, lists and subheadings inside section HTML
    are rendered inline (in order) rather than grouped at the end."""

    @staticmethod
    def _make_base64_png() -> str:
        """Return a base64 data URI for a minimal 1x1 white PNG."""
        import struct
        import zlib
        sig = b'\x89PNG\r\n\x1a\n'
        ihdr_data = struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)
        ihdr_crc = zlib.crc32(b'IHDR' + ihdr_data) & 0xffffffff
        ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data + struct.pack('>I', ihdr_crc)
        raw = zlib.compress(b'\x00\xff\xff\xff')
        idat_crc = zlib.crc32(b'IDAT' + raw) & 0xffffffff
        idat = struct.pack('>I', len(raw)) + b'IDAT' + raw + struct.pack('>I', idat_crc)
        iend_crc = zlib.crc32(b'IEND') & 0xffffffff
        iend = struct.pack('>I', 0) + b'IEND' + struct.pack('>I', iend_crc)
        import base64 as b64
        png_bytes = sig + ihdr + idat + iend
        encoded = b64.b64encode(png_bytes).decode()
        return f"data:image/png;base64,{encoded}"

    def test_inline_base64_image_in_section(self, tmp_dir):
        """A base64 <figure> inside a section should render inline, not at the end."""
        data_uri = self._make_base64_png()
        html = (
            '<p>Before the chart.</p>'
            f'<figure><img src="{data_uri}" alt="Inline Chart"/>'
            '<figcaption>Figure 1</figcaption></figure>'
            '<p>After the chart.</p>'
        )
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Inline Image Test",
            analysis=html,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "inline_img.docx")
        doc = Document(str(path))

        # Collect paragraph texts and check ordering
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        before_idx = next(i for i, t in enumerate(texts) if "Before the chart" in t)
        caption_idx = next(i for i, t in enumerate(texts) if "Figure 1" in t)
        after_idx = next(i for i, t in enumerate(texts) if "After the chart" in t)
        assert before_idx < caption_idx < after_idx

    def test_inline_table_in_section(self, tmp_dir):
        """An HTML table inside section content should render inline."""
        html = (
            '<p>Intro text.</p>'
            '<table><tr><th>Region</th><th>Value</th></tr>'
            '<tr><td>North</td><td>100</td></tr></table>'
            '<p>Conclusion text.</p>'
        )
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Inline Table Test",
            analysis=html,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "inline_table.docx")
        doc = Document(str(path))

        # The table should exist and contain Region header
        found = False
        for table in doc.tables:
            cell_texts = [cell.text for row in table.rows for cell in row.cells]
            if "Region" in cell_texts and "North" in cell_texts:
                found = True
                break
        assert found

        # Text before and after should exist
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Intro text" in t for t in texts)
        assert any("Conclusion text" in t for t in texts)

    def test_inline_subheadings_and_lists(self, tmp_dir):
        """<h3> and <ul> inside section HTML should render inline."""
        html = (
            '<p>Opening paragraph.</p>'
            '<h3>Sub-section Title</h3>'
            '<ul><li>Item A</li><li>Item B</li></ul>'
            '<p>Closing paragraph.</p>'
        )
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Subheadings Test",
            analysis=html,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "inline_sub.docx")
        doc = Document(str(path))

        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Sub-section Title" in t for t in texts)
        assert any("Item A" in t for t in texts)
        assert any("Item B" in t for t in texts)

        # Verify ordering: opening < sub-section title < items < closing
        opening_idx = next(i for i, t in enumerate(texts) if "Opening paragraph" in t)
        heading_idx = next(i for i, t in enumerate(texts) if "Sub-section Title" in t)
        closing_idx = next(i for i, t in enumerate(texts) if "Closing paragraph" in t)
        assert opening_idx < heading_idx < closing_idx

    def test_plain_html_still_works(self, tmp_dir):
        """Backward compat: simple <p> with bold/italic renders correctly."""
        html = '<p>This is <strong>bold</strong> and <em>italic</em> text.</p>'
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Plain HTML Test",
            analysis=html,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "plain_html.docx")
        doc = Document(str(path))

        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("bold" in t and "italic" in t for t in texts)

    def test_figures_list_still_appended(self, tmp_dir, sample_figures):
        """The figures= parameter should still append figures at the end."""
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Figures Param Test",
            analysis="<p>Some analysis text.</p>",
            figures=sample_figures,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "figures_param.docx")
        doc = Document(str(path))

        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Some analysis text" in t for t in texts)
        assert any("Test Chart" in t for t in texts)


# ---------------------------------------------------------------------------
# Pagination tests
# ---------------------------------------------------------------------------

class TestPagination:
    """Verify pagination control properties (cantSplit, tblHeader,
    keep_with_next, page_break_before, widow_control)."""

    def test_widow_control_on_normal_style(self, tmp_dir):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(title="Widow Test", show_cover=False)
        path = gen.save(tmp_dir / "widow.docx")
        doc = Document(str(path))
        assert doc.styles["Normal"].paragraph_format.widow_control is True

    def test_h2_sections_have_page_break_before(self, tmp_dir):
        """First section has no page break; subsequent sections do."""
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="PageBreak Test",
            executive_summary="<p>Summary</p>",
            methodology="<p>Method</p>",
            analysis="<p>Analysis</p>",
            show_cover=False,
        )
        path = gen.save(tmp_dir / "pagebreak.docx")
        doc = Document(str(path))
        h2_paras = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        assert len(h2_paras) >= 2
        # First H2 section: no page break before
        assert h2_paras[0].paragraph_format.page_break_before in (None, False)
        # Subsequent H2 sections: page break before
        for h2 in h2_paras[1:]:
            assert h2.paragraph_format.page_break_before is True

    def test_headings_have_keep_with_next(self, tmp_dir):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="KeepNext Test",
            analysis="<h3>Sub-section</h3><p>Content</p>",
            show_cover=False,
        )
        path = gen.save(tmp_dir / "keepnext.docx")
        doc = Document(str(path))
        # H3 rendered inline has explicit keep_with_next on the paragraph
        h3_paras = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
        assert len(h3_paras) >= 1
        for h in h3_paras:
            assert h.paragraph_format.keep_with_next is True
        # Heading styles (1, 2, 3) have keep_with_next set at style level
        for level in (1, 2, 3):
            style = doc.styles[f"Heading {level}"]
            assert style.paragraph_format.keep_with_next is True

    def test_table_rows_cant_split(self, tmp_dir, sample_tables):
        from docx.oxml.ns import qn as _qn
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="CantSplit Test",
            data_tables=sample_tables,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "cantsplit.docx")
        doc = Document(str(path))
        # Find the data table (has Region header)
        for table in doc.tables:
            cell_texts = [c.text for r in table.rows for c in r.cells]
            if "Region" in cell_texts:
                for row in table.rows:
                    trPr = row._tr.find(_qn("w:trPr"))
                    assert trPr is not None
                    cant_split = trPr.find(_qn("w:cantSplit"))
                    assert cant_split is not None
                    assert cant_split.get(_qn("w:val")) == "true"
                break

    def test_table_header_repeats(self, tmp_dir, sample_tables):
        from docx.oxml.ns import qn as _qn
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="TblHeader Test",
            data_tables=sample_tables,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "tblheader.docx")
        doc = Document(str(path))
        for table in doc.tables:
            cell_texts = [c.text for r in table.rows for c in r.cells]
            if "Region" in cell_texts:
                trPr = table.rows[0]._tr.find(_qn("w:trPr"))
                assert trPr is not None
                tbl_header = trPr.find(_qn("w:tblHeader"))
                assert tbl_header is not None
                assert tbl_header.get(_qn("w:val")) == "true"
                break

    def test_figure_keeps_with_caption(self, tmp_dir, sample_figures):
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="FigKeep Test",
            figures=sample_figures,
            show_cover=False,
        )
        path = gen.save(tmp_dir / "figkeep.docx")
        doc = Document(str(path))
        # Find the paragraph just before the caption "Test Chart"
        for i, p in enumerate(doc.paragraphs):
            if "Test Chart" in p.text and p.style.name == "Caption":
                # Previous paragraph should have keep_with_next
                prev = doc.paragraphs[i - 1]
                assert prev.paragraph_format.keep_with_next is True
                break

    def test_callout_row_cant_split(self, tmp_dir, sample_callouts):
        from docx.oxml.ns import qn as _qn
        gen = DOCXGenerator(style="corporate")
        gen.render_scaffold(
            title="Callout CantSplit",
            callouts=[sample_callouts[0]],
            show_cover=False,
        )
        path = gen.save(tmp_dir / "callout_cs.docx")
        doc = Document(str(path))
        # The callout table is the last table
        table = doc.tables[-1]
        assert len(table.rows) == 1
        trPr = table.rows[0]._tr.find(_qn("w:trPr"))
        assert trPr is not None
        cant_split = trPr.find(_qn("w:cantSplit"))
        assert cant_split is not None
        assert cant_split.get(_qn("w:val")) == "true"


class TestSaveErrors:
    def test_save_without_render_raises(self, tmp_dir):
        gen = DOCXGenerator(style="corporate")
        with pytest.raises(RuntimeError, match="render_scaffold"):
            gen.save(tmp_dir / "error.docx")

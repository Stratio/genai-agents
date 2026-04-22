"""
quality_report_generator.py
---------------------------
Generates data quality coverage reports in PDF, DOCX or Markdown format.

Usage:
    python3 scripts/quality_report_generator.py --format pdf --output output/report.pdf --input-json '{"title": ...}'
    python3 scripts/quality_report_generator.py --format docx --output output/report.docx --input-file data.json
    python3 scripts/quality_report_generator.py --format md --output output/report.md --input-json '{"title": ...}'

Input JSON format:
    {
        "title": "Data Quality Coverage Report",
        "domain": "semantic_financial",
        "scope": "Full domain",
        "generated_at": "2025-02-26",
        "summary": {
            "tables_analyzed": 5,
            "rules_total": 12,
            "rules_ok": 8,
            "rules_ko": 2,
            "rules_warning": 1,
            "rules_not_executed": 1,
            "coverage_estimate": "62%",
            "gaps_critical": 4,
            "gaps_moderate": 3,
            "gaps_low": 2,
            "rules_created_this_session": 3
        },
        "tables": [
            {
                "name": "account",
                "coverage_estimate": "75%",
                "completeness": "OK",
                "uniqueness": "Gap",
                "validity": "Parcial",
                "consistency": "N/A",
                "rules": [
                    {
                        "name": "dq-account-completeness-id",
                        "dimension": "completeness",
                        "status": "OK",
                        "pass_pct": 100.0,
                        "description": "The id field is never null"
                    }
                ],
                "gaps": [
                    {
                        "column": "account_id",
                        "dimension": "uniqueness",
                        "priority": "CRITICO",
                        "description": "Primary key without uniqueness verification"
                    }
                ]
            }
        ],
        "rules_created": [
            {
                "name": "dq-account-uniqueness-id",
                "table": "account",
                "dimension": "uniqueness",
                "status": "created"
            }
        ],
        "recommendations": [
            "Resolve the 2 rules in KO status before creating new ones",
            "Cover critical primary key gaps with uniqueness rules"
        ]
    }
"""

import argparse
import json
import sys
import os
from datetime import date

# Local import — scripts/ is on sys.path when invoked directly, and the file
# sits alongside this module once packaged inside the agent.
_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)
from i18n import get_labels  # noqa: E402


# ---------------------------------------------------------------------------
# Common helpers
# ---------------------------------------------------------------------------

STATUS_ICONS = {
    "OK": "[OK]",
    "KO": "[KO]",
    "WARNING": "[WARN]",
    "N/A": "[N/A]",
    "Gap": "[GAP]",
    "Parcial": "[PARTIAL]",
    "created": "[CREATED]",
    "error": "[ERROR]",
}

PRIORITY_LABELS = {"CRITICO": "CRITICO", "ALTO": "ALTO", "MEDIO": "MEDIO", "BAJO": "BAJO"}


def load_input(input_json: str | None, input_file: str | None) -> dict:
    if input_json:
        return json.loads(input_json)
    if input_file:
        with open(input_file, encoding="utf-8") as f:
            return json.load(f)
    # Try reading from stdin
    raw = sys.stdin.read().strip()
    if raw:
        return json.loads(raw)
    raise ValueError("No input data provided (--input-json, --input-file or stdin)")


def resolve_labels_for_data(
    data: dict,
    cli_lang: str | None = None,
    cli_labels: dict | None = None,
) -> dict[str, str]:
    """Resolve the effective label dict for a report invocation.

    Priority (highest wins):
      1. `cli_labels` (from `--labels-json`)
      2. `data["labels"]` embedded in the JSON input
      3. `cli_lang` (from `--lang`)
      4. `data["lang"]` embedded in the JSON input
      5. `.agent_lang` file at the packaged agent root
      6. English fallback
    """
    effective_lang = cli_lang or data.get("lang")
    merged_overrides: dict[str, str] = {}
    if isinstance(data.get("labels"), dict):
        merged_overrides.update(data["labels"])
    if cli_labels:
        merged_overrides.update(cli_labels)
    return get_labels(lang=effective_lang, overrides=merged_overrides or None)


def ensure_output_dir(output_path: str) -> None:
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)


# ---------------------------------------------------------------------------
# Markdown generation
# ---------------------------------------------------------------------------

def build_markdown(data: dict, labels: dict[str, str] | None = None) -> str:
    if labels is None:
        labels = resolve_labels_for_data(data)
    s = data.get("summary", {})
    lines = []

    # Header
    lines.append(f"# {data.get('title', labels['quality.default_title'])}")
    lines.append("")
    lines.append(f"**{labels['quality.meta.collection']}**: {data.get('domain', '-')}")
    lines.append(f"**{labels['quality.meta.scope']}**: {data.get('scope', '-')}")
    lines.append(f"**{labels['quality.meta.generation_date']}**: {data.get('generated_at', str(date.today()))}")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Executive summary
    lines.append(f"## {labels['quality.executive_summary']}")
    lines.append("")
    lines.append(f"| {labels['quality.col.metric']} | {labels['quality.col.value']} |")
    lines.append(f"|--------|-------|")
    lines.append(f"| {labels['quality.summary.tables_analyzed']} | {s.get('tables_analyzed', '-')} |")
    lines.append(f"| {labels['quality.summary.existing_rules']} | {s.get('rules_total', '-')} |")
    lines.append(f"| {labels['quality.summary.status_ok']} | {s.get('rules_ok', '-')} |")
    lines.append(f"| {labels['quality.summary.status_ko']} | {s.get('rules_ko', '-')} |")
    lines.append(f"| {labels['quality.summary.status_warning']} | {s.get('rules_warning', '-')} |")
    lines.append(f"| {labels['quality.summary.not_executed']} | {s.get('rules_not_executed', '-')} |")
    lines.append(f"| {labels['quality.summary.estimated_coverage']} | **{s.get('coverage_estimate', '-')}** |")
    lines.append(f"| {labels['quality.summary.gaps_critical']} | {s.get('gaps_critical', '-')} |")
    lines.append(f"| {labels['quality.summary.gaps_moderate']} | {s.get('gaps_moderate', '-')} |")
    lines.append(f"| {labels['quality.summary.gaps_low']} | {s.get('gaps_low', '-')} |")
    if s.get("rules_created_this_session"):
        lines.append(f"| {labels['quality.summary.rules_created_this_session']} | {s['rules_created_this_session']} |")
    lines.append("")

    # Coverage by table
    tables = data.get("tables", [])
    if tables:
        lines.append(f"## {labels['quality.coverage_by_table']}")
        lines.append("")
        lines.append(
            f"| {labels['quality.col.table']} "
            f"| {labels['quality.col.completeness']} "
            f"| {labels['quality.col.uniqueness']} "
            f"| {labels['quality.col.validity']} "
            f"| {labels['quality.col.consistency']} "
            f"| {labels['quality.col.coverage']} |"
        )
        lines.append("|-------|-------------|------------|----------|-------------|----------|")
        for t in tables:
            lines.append(
                f"| {t['name']} "
                f"| {t.get('completeness', '-')} "
                f"| {t.get('uniqueness', '-')} "
                f"| {t.get('validity', '-')} "
                f"| {t.get('consistency', '-')} "
                f"| {t.get('coverage_estimate', '-')} |"
            )
        lines.append("")

        # Existing rules detail
        lines.append(f"## {labels['quality.existing_rules_detail']}")
        lines.append("")
        for t in tables:
            rules = t.get("rules", [])
            if not rules:
                continue
            lines.append(f"### {t['name']}")
            lines.append("")
            lines.append(
                f"| {labels['quality.col.rule']} "
                f"| {labels['quality.col.dimension']} "
                f"| {labels['quality.col.status']} "
                f"| {labels['quality.col.pass_pct']} "
                f"| {labels['quality.col.description']} |"
            )
            lines.append("|------|-----------|--------|--------|-------------|")
            for r in rules:
                pass_pct = f"{r['pass_pct']:.1f}%" if r.get("pass_pct") is not None else "-"
                status = r.get("status", "-")
                status_marker = "**" if status in ("KO", "WARNING") else ""
                lines.append(
                    f"| {status_marker}{r['name']}{status_marker} "
                    f"| {r.get('dimension', '-')} "
                    f"| {status_marker}{status}{status_marker} "
                    f"| {pass_pct} "
                    f"| {r.get('description', '-')} |"
                )
            lines.append("")

        # Identified gaps
        all_gaps = []
        for t in tables:
            for g in t.get("gaps", []):
                all_gaps.append({**g, "table": t["name"]})

        if all_gaps:
            lines.append(f"## {labels['quality.identified_gaps']}")
            lines.append("")
            # Order: CRITICO > ALTO > MEDIO > BAJO
            priority_order = {"CRITICO": 0, "ALTO": 1, "MEDIO": 2, "BAJO": 3}
            all_gaps.sort(key=lambda x: priority_order.get(x.get("priority", "BAJO"), 99))
            lines.append(
                f"| {labels['quality.col.priority']} "
                f"| {labels['quality.col.table']} "
                f"| {labels['quality.col.column']} "
                f"| {labels['quality.col.dimension']} "
                f"| {labels['quality.col.description']} |"
            )
            lines.append("|----------|-------|--------|-----------|-------------|")
            for g in all_gaps:
                lines.append(
                    f"| {g.get('priority', '-')} "
                    f"| {g.get('table', '-')} "
                    f"| {g.get('column', '-')} "
                    f"| {g.get('dimension', '-')} "
                    f"| {g.get('description', '-')} |"
                )
            lines.append("")

    # Rules created this session
    rules_created = data.get("rules_created", [])
    if rules_created:
        lines.append(f"## {labels['quality.rules_created_session']}")
        lines.append("")
        lines.append(
            f"| {labels['quality.col.rule']} "
            f"| {labels['quality.col.table']} "
            f"| {labels['quality.col.dimension']} "
            f"| {labels['quality.col.status']} |"
        )
        lines.append("|------|-------|-----------|--------|")
        for r in rules_created:
            lines.append(
                f"| {r['name']} "
                f"| {r.get('table', '-')} "
                f"| {r.get('dimension', '-')} "
                f"| {r.get('status', '-')} |"
            )
        lines.append("")

    # Recommendations
    recommendations = data.get("recommendations", [])
    if recommendations:
        lines.append(f"## {labels['quality.recommendations']}")
        lines.append("")
        for i, rec in enumerate(recommendations, 1):
            lines.append(f"{i}. {rec}")
        lines.append("")

    lines.append("---")
    lines.append(f"*{labels['quality.footer']}*")
    lines.append("")

    return "\n".join(lines)


def generate_markdown(data: dict, output_path: str,
                      labels: dict[str, str] | None = None) -> None:
    content = build_markdown(data, labels=labels)
    ensure_output_dir(output_path)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"[OK] Markdown generated: {output_path}")


# ---------------------------------------------------------------------------
# PDF generation via WeasyPrint
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Visual tones — optional stylistic direction for the report
# ---------------------------------------------------------------------------
#
# Each tone is a small CSS override that changes type pairing and palette
# without touching the HTML body. See `skill-guides/visual-craftsmanship.md`
# for the shared aesthetic principles.

TONES = {
    "default": {
        "body_font": "Arial, sans-serif",
        "display_font": "Arial, sans-serif",
        "mono_font": "monospace",
        "ink": "#1a1a2e",
        "accent": "#0f3460",
        "accent_deep": "#16213e",
        "rule_subtle": "#dddddd",
        "surface_soft": "#f0f4ff",
        "surface_alt": "#f8f9fa",
    },
    "technical-minimal": {
        "body_font": "'IBM Plex Serif', Georgia, serif",
        "display_font": "'IBM Plex Sans', Helvetica, sans-serif",
        "mono_font": "'IBM Plex Mono', Consolas, monospace",
        "ink": "#16191f",
        "accent": "#0a4b6e",
        "accent_deep": "#062b41",
        "rule_subtle": "#dadfe4",
        "surface_soft": "#eef3f7",
        "surface_alt": "#f5f7f9",
    },
    "executive-editorial": {
        "body_font": "'Crimson Pro', Georgia, serif",
        "display_font": "'Instrument Serif', Georgia, serif",
        "mono_font": "'JetBrains Mono', Consolas, monospace",
        "ink": "#201a16",
        "accent": "#8a3324",
        "accent_deep": "#571c14",
        "rule_subtle": "#dccfbf",
        "surface_soft": "#f5ecdf",
        "surface_alt": "#faf4ea",
    },
    "forensic": {
        "body_font": "'IBM Plex Mono', Consolas, monospace",
        "display_font": "'IBM Plex Serif', Georgia, serif",
        "mono_font": "'IBM Plex Mono', Consolas, monospace",
        "ink": "#121212",
        "accent": "#5a1c1c",
        "accent_deep": "#2f0f0f",
        "rule_subtle": "#cccccc",
        "surface_soft": "#f2ebe8",
        "surface_alt": "#ecebe8",
    },
}


def _tone_palette(tone: str | None) -> dict:
    """Return the palette for a tone name, falling back to `default`."""
    if not tone:
        return TONES["default"]
    return TONES.get(tone, TONES["default"])


HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="{{ lang }}">
<head>
<meta charset="UTF-8">
<style>
  :root {
    --body-font: {{ body_font }};
    --display-font: {{ display_font }};
    --mono-font: {{ mono_font }};
    --ink: {{ ink }};
    --accent: {{ accent }};
    --accent-deep: {{ accent_deep }};
    --rule-subtle: {{ rule_subtle }};
    --surface-soft: {{ surface_soft }};
    --surface-alt: {{ surface_alt }};
  }
  body { font-family: var(--body-font); font-size: 11pt; color: var(--ink); margin: 40px; }
  h1 { font-family: var(--display-font); color: var(--accent-deep); border-bottom: 3px solid var(--accent); padding-bottom: 8px; font-size: 20pt; }
  h2 { font-family: var(--display-font); color: var(--accent-deep); border-left: 4px solid var(--accent); padding-left: 10px; font-size: 14pt; margin-top: 30px; }
  h3 { font-family: var(--display-font); color: var(--accent); font-size: 12pt; margin-top: 20px; }
  .meta { color: #555; font-size: 10pt; margin-bottom: 20px; }
  .meta span { margin-right: 20px; }
  table { width: 100%; border-collapse: collapse; margin: 10px 0 20px 0; font-size: 10pt; }
  th { background-color: var(--accent); color: white; padding: 7px 10px; text-align: left; font-family: var(--display-font); }
  td { padding: 6px 10px; border-bottom: 1px solid var(--rule-subtle); }
  th, td { word-wrap: break-word; overflow-wrap: anywhere; hyphens: auto; vertical-align: top; }
  td.num { font-family: var(--mono-font); font-variant-numeric: tabular-nums; }
  tr:nth-child(even) { background-color: var(--surface-alt); }
  .fixed-table { table-layout: fixed; }
  .fixed-table td, .fixed-table th { font-size: 9.5pt; }
  .status-ok { color: #28a745; font-weight: bold; }
  .status-ko { color: #dc3545; font-weight: bold; }
  .status-warning { color: #ffc107; font-weight: bold; }
  .status-gap { color: #6c757d; }
  .priority-critico { color: #dc3545; font-weight: bold; }
  .priority-alto { color: #fd7e14; font-weight: bold; }
  .priority-medio { color: #ffc107; }
  .priority-bajo { color: #6c757d; }
  .summary-box { background: var(--surface-soft); border: 1px solid var(--rule-subtle); border-radius: 6px;
                 padding: 15px 20px; margin: 20px 0; }
  .summary-box table { margin: 0; }
  .summary-box th { background-color: var(--accent); }
  .rec-list { padding-left: 20px; }
  .rec-list li { margin-bottom: 6px; }
  .footer { text-align: center; font-size: 9pt; color: #888; border-top: 1px solid var(--rule-subtle);
            margin-top: 40px; padding-top: 10px; }
  @page { margin: 2cm; }
</style>
</head>
<body>
{{ content_html }}
<div class="footer">{{ footer_text }} | {{ generated_at }}</div>
</body>
</html>
"""


def _apply_tone(html_template: str, tone: str | None) -> str:
    """Substitute the tone palette placeholders in the HTML template."""
    palette = _tone_palette(tone)
    result = html_template
    for key, value in palette.items():
        result = result.replace("{{ " + key + " }}", value)
    return result


def _status_class(status: str) -> str:
    mapping = {"OK": "status-ok", "KO": "status-ko", "WARNING": "status-warning",
               "Gap": "status-gap", "Parcial": "status-gap"}
    return mapping.get(status, "")


def _priority_class(priority: str) -> str:
    mapping = {"CRITICO": "priority-critico", "ALTO": "priority-alto",
               "MEDIO": "priority-medio", "BAJO": "priority-bajo"}
    return mapping.get(priority, "")


def build_html_body(data: dict, labels: dict[str, str] | None = None) -> str:
    if labels is None:
        labels = resolve_labels_for_data(data)
    s = data.get("summary", {})
    parts = []

    # Header
    parts.append(f"<h1>{data.get('title', labels['quality.default_title'])}</h1>")
    parts.append(
        f"<div class='meta'>"
        f"<span><b>{labels['quality.meta.domain']}:</b> {data.get('domain', '-')}</span>"
        f"<span><b>{labels['quality.meta.scope']}:</b> {data.get('scope', '-')}</span>"
        f"<span><b>{labels['quality.meta.date']}:</b> {data.get('generated_at', str(date.today()))}</span>"
        f"</div>"
    )

    # Executive summary
    parts.append(f"<h2>{labels['quality.executive_summary']}</h2>")
    parts.append("<div class='summary-box'>")
    parts.append(f"<table><tr><th>{labels['quality.col.metric']}</th><th>{labels['quality.col.value']}</th></tr>")
    summary_rows = [
        (labels['quality.summary.tables_analyzed'], s.get("tables_analyzed", "-")),
        (labels['quality.summary.existing_rules'], s.get("rules_total", "-")),
        (labels['quality.summary.status_breakdown'], f"{s.get('rules_ok',0)} / {s.get('rules_ko',0)} / {s.get('rules_warning',0)}"),
        (labels['quality.summary.not_executed'], s.get("rules_not_executed", "-")),
        (labels['quality.summary.estimated_coverage'], f"<b>{s.get('coverage_estimate', '-')}</b>"),
        (labels['quality.summary.gaps_breakdown'],
         f"{s.get('gaps_critical',0)} / {s.get('gaps_moderate',0)} / {s.get('gaps_low',0)}"),
    ]
    if s.get("rules_created_this_session"):
        summary_rows.append((labels['quality.summary.rules_created_this_session'], s["rules_created_this_session"]))
    for label, value in summary_rows:
        parts.append(f"<tr><td>{label}</td><td>{value}</td></tr>")
    parts.append("</table></div>")

    # Coverage by table
    tables = data.get("tables", [])
    if tables:
        parts.append(f"<h2>{labels['quality.coverage_by_table']}</h2>")
        parts.append(
            f"<table><tr>"
            f"<th>{labels['quality.col.table']}</th>"
            f"<th>{labels['quality.col.completeness']}</th>"
            f"<th>{labels['quality.col.uniqueness']}</th>"
            f"<th>{labels['quality.col.validity']}</th>"
            f"<th>{labels['quality.col.consistency']}</th>"
            f"<th>{labels['quality.col.coverage']}</th></tr>"
        )
        for t in tables:
            def cell(val):
                cls = _status_class(val or "-")
                return f"<td class='{cls}'>{val or '-'}</td>"
            parts.append(
                f"<tr><td><b>{t['name']}</b></td>"
                f"{cell(t.get('completeness'))}"
                f"{cell(t.get('uniqueness'))}"
                f"{cell(t.get('validity'))}"
                f"{cell(t.get('consistency'))}"
                f"<td><b>{t.get('coverage_estimate', '-')}</b></td></tr>"
            )
        parts.append("</table>")

        # Rules detail
        parts.append(f"<h2>{labels['quality.existing_rules_detail']}</h2>")
        for t in tables:
            rules = t.get("rules", [])
            if not rules:
                continue
            parts.append(f"<h3>{t['name']}</h3>")
            parts.append(
                f"<table class='fixed-table'>"
                f"<colgroup>"
                f"<col style='width:26%'><col style='width:14%'>"
                f"<col style='width:10%'><col style='width:10%'>"
                f"<col style='width:40%'>"
                f"</colgroup>"
                f"<tr>"
                f"<th>{labels['quality.col.rule']}</th>"
                f"<th>{labels['quality.col.dimension']}</th>"
                f"<th>{labels['quality.col.status']}</th>"
                f"<th>{labels['quality.col.pass_pct']}</th>"
                f"<th>{labels['quality.col.description']}</th></tr>"
            )
            for r in rules:
                status = r.get("status", "-")
                cls = _status_class(status)
                pass_pct = f"{r['pass_pct']:.1f}%" if r.get("pass_pct") is not None else "-"
                parts.append(
                    f"<tr><td>{r['name']}</td><td>{r.get('dimension','-')}</td>"
                    f"<td class='{cls}'>{status}</td><td>{pass_pct}</td>"
                    f"<td>{r.get('description','-')}</td></tr>"
                )
            parts.append("</table>")

        # Gaps
        all_gaps = []
        for t in tables:
            for g in t.get("gaps", []):
                all_gaps.append({**g, "table": t["name"]})

        if all_gaps:
            priority_order = {"CRITICO": 0, "ALTO": 1, "MEDIO": 2, "BAJO": 3}
            all_gaps.sort(key=lambda x: priority_order.get(x.get("priority", "BAJO"), 99))
            parts.append(f"<h2>{labels['quality.identified_gaps']}</h2>")
            parts.append(
                f"<table class='fixed-table'>"
                f"<colgroup>"
                f"<col style='width:10%'><col style='width:18%'>"
                f"<col style='width:18%'><col style='width:14%'>"
                f"<col style='width:40%'>"
                f"</colgroup>"
                f"<tr>"
                f"<th>{labels['quality.col.priority']}</th>"
                f"<th>{labels['quality.col.table']}</th>"
                f"<th>{labels['quality.col.column']}</th>"
                f"<th>{labels['quality.col.dimension']}</th>"
                f"<th>{labels['quality.col.description']}</th></tr>"
            )
            for g in all_gaps:
                pri = g.get("priority", "-")
                cls = _priority_class(pri)
                parts.append(
                    f"<tr><td class='{cls}'>{pri}</td><td>{g.get('table','-')}</td>"
                    f"<td><b>{g.get('column','-')}</b></td><td>{g.get('dimension','-')}</td>"
                    f"<td>{g.get('description','-')}</td></tr>"
                )
            parts.append("</table>")

    # Rules created
    rules_created = data.get("rules_created", [])
    if rules_created:
        parts.append(f"<h2>{labels['quality.rules_created_session']}</h2>")
        parts.append(
            f"<table><tr>"
            f"<th>{labels['quality.col.rule']}</th>"
            f"<th>{labels['quality.col.table']}</th>"
            f"<th>{labels['quality.col.dimension']}</th>"
            f"<th>{labels['quality.col.status']}</th></tr>"
        )
        for r in rules_created:
            status = r.get("status", "-")
            cls = "status-ok" if status == "created" else "status-ko"
            parts.append(
                f"<tr><td>{r['name']}</td><td>{r.get('table','-')}</td>"
                f"<td>{r.get('dimension','-')}</td><td class='{cls}'>{status}</td></tr>"
            )
        parts.append("</table>")

    # Recommendations
    recommendations = data.get("recommendations", [])
    if recommendations:
        parts.append(f"<h2>{labels['quality.recommendations']}</h2>")
        parts.append("<ol class='rec-list'>")
        for rec in recommendations:
            parts.append(f"<li>{rec}</li>")
        parts.append("</ol>")

    return "\n".join(parts)


def generate_pdf(data: dict, output_path: str,
                 labels: dict[str, str] | None = None,
                 tone: str | None = None) -> None:
    try:
        from weasyprint import HTML as WeasyprintHTML
    except ImportError:
        print("[ERROR] weasyprint is not installed. Run: pip install weasyprint", file=sys.stderr)
        sys.exit(1)

    if labels is None:
        labels = resolve_labels_for_data(data)

    html_body = build_html_body(data, labels=labels)
    generated_at = data.get("generated_at", str(date.today()))
    html_full = (
        _apply_tone(HTML_TEMPLATE, tone)
        .replace("{{ content_html }}", html_body)
        .replace("{{ generated_at }}", generated_at)
        .replace("{{ lang }}", labels["html.lang_attr"])
        .replace("{{ footer_text }}", labels["quality.footer"])
    )

    ensure_output_dir(output_path)
    WeasyprintHTML(string=html_full).write_pdf(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f"[OK] PDF generated: {output_path} ({size_kb} KB)")


# ---------------------------------------------------------------------------
# DOCX generation via python-docx
# ---------------------------------------------------------------------------

def _hex_to_rgbcolor(hex_str: str, rgb_color_cls):
    """Convert '#rrggbb' to a python-docx RGBColor instance."""
    h = hex_str.lstrip("#")
    return rgb_color_cls(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def generate_docx(data: dict, output_path: str,
                  labels: dict[str, str] | None = None,
                  tone: str | None = None) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[ERROR] python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    if labels is None:
        labels = resolve_labels_for_data(data)

    palette = _tone_palette(tone)

    doc = Document()

    # Basic styles — accent drawn from the chosen tone, status colours kept fixed
    BLUE = _hex_to_rgbcolor(palette["accent_deep"], RGBColor)
    RED = RGBColor(220, 53, 69)
    GREEN = RGBColor(40, 167, 69)
    ORANGE = RGBColor(253, 126, 20)
    AMBER = RGBColor(255, 193, 7)
    GRAY = RGBColor(108, 117, 125)

    # Status → colour map used to tint OK / KO / WARNING / Gap / Parcial values
    # in the coverage table. Priorities (CRITICO / ALTO / MEDIO / BAJO) use their
    # own map. Applied in-place to the first run of the target cell.
    _STATUS_CELL_COLOURS = {
        "OK": GREEN,
        "KO": RED,
        "WARNING": ORANGE,
        "Parcial": ORANGE,
        "Gap": GRAY,
    }
    _PRIORITY_CELL_COLOURS = {
        "CRITICO": RED,
        "ALTO": ORANGE,
        "MEDIO": AMBER,
        "BAJO": GRAY,
    }

    def _tint_cell(cell, value: str, mapping: dict) -> None:
        """If ``value`` has an entry in ``mapping``, make the cell's first run
        bold and coloured accordingly. No-op otherwise."""
        colour = mapping.get(value)
        if colour is None:
            return
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = colour

    s = data.get("summary", {})

    # Title
    title = doc.add_heading(data.get("title", labels["quality.default_title"]), level=0)
    title.runs[0].font.color.rgb = BLUE

    # Meta
    meta = doc.add_paragraph()
    meta.add_run(f"{labels['quality.meta.domain']}: ").bold = True
    meta.add_run(f"{data.get('domain', '-')}   ")
    meta.add_run(f"{labels['quality.meta.scope']}: ").bold = True
    meta.add_run(f"{data.get('scope', '-')}   ")
    meta.add_run(f"{labels['quality.meta.date']}: ").bold = True
    meta.add_run(data.get("generated_at", str(date.today())))

    doc.add_heading(labels['quality.executive_summary'], level=1)

    # Summary table
    summary_rows = [
        (labels['quality.summary.tables_analyzed'], str(s.get("tables_analyzed", "-"))),
        (labels['quality.summary.existing_rules'], str(s.get("rules_total", "-"))),
        (labels['quality.summary.status_breakdown'],
         f"{s.get('rules_ok',0)} / {s.get('rules_ko',0)} / {s.get('rules_warning',0)}"),
        (labels['quality.summary.not_executed'], str(s.get("rules_not_executed", "-"))),
        (labels['quality.summary.estimated_coverage'], str(s.get("coverage_estimate", "-"))),
        (labels['quality.summary.gaps_breakdown'],
         f"{s.get('gaps_critical',0)} / {s.get('gaps_moderate',0)} / {s.get('gaps_low',0)}"),
    ]
    if s.get("rules_created_this_session"):
        summary_rows.append((labels['quality.summary.rules_created_this_session'], str(s["rules_created_this_session"])))

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = labels['quality.col.metric']
    hdr[1].text = labels['quality.col.value']
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    for label, value in summary_rows:
        row = tbl.add_row().cells
        row[0].text = label
        row[1].text = value

    # Coverage by table
    tables = data.get("tables", [])
    if tables:
        doc.add_heading(labels['quality.coverage_by_table'], level=1)
        tbl2 = doc.add_table(rows=1, cols=6)
        tbl2.style = "Table Grid"
        headers = [
            labels['quality.col.table'],
            labels['quality.col.completeness'],
            labels['quality.col.uniqueness'],
            labels['quality.col.validity'],
            labels['quality.col.consistency'],
            labels['quality.col.coverage'],
        ]
        for i, h in enumerate(headers):
            tbl2.rows[0].cells[i].text = h
            tbl2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for t in tables:
            row = tbl2.add_row().cells
            vals = [t["name"], t.get("completeness", "-"), t.get("uniqueness", "-"),
                    t.get("validity", "-"), t.get("consistency", "-"), t.get("coverage_estimate", "-")]
            for i, v in enumerate(vals):
                row[i].text = v
            # Tint the four dimension-status columns (completeness, uniqueness,
            # validity, consistency). The coverage column is a percentage and
            # stays neutral.
            for i in range(1, 5):
                _tint_cell(row[i], vals[i], _STATUS_CELL_COLOURS)

        # Gaps
        all_gaps = []
        for t in tables:
            for g in t.get("gaps", []):
                all_gaps.append({**g, "table": t["name"]})
        if all_gaps:
            priority_order = {"CRITICO": 0, "ALTO": 1, "MEDIO": 2, "BAJO": 3}
            all_gaps.sort(key=lambda x: priority_order.get(x.get("priority", "BAJO"), 99))
            doc.add_heading(labels['quality.identified_gaps'], level=1)
            tbl3 = doc.add_table(rows=1, cols=5)
            tbl3.style = "Table Grid"
            gap_headers = [
                labels['quality.col.priority'],
                labels['quality.col.table'],
                labels['quality.col.column'],
                labels['quality.col.dimension'],
                labels['quality.col.description'],
            ]
            for i, h in enumerate(gap_headers):
                tbl3.rows[0].cells[i].text = h
                tbl3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for g in all_gaps:
                row = tbl3.add_row().cells
                vals = [g.get("priority", "-"), g.get("table", "-"),
                        g.get("column", "-"), g.get("dimension", "-"), g.get("description", "-")]
                for i, v in enumerate(vals):
                    row[i].text = v
                # Tint the priority column (col 0) so CRITICO/ALTO stand out.
                _tint_cell(row[0], vals[0], _PRIORITY_CELL_COLOURS)

    # Rules created
    rules_created = data.get("rules_created", [])
    if rules_created:
        doc.add_heading(labels['quality.rules_created_session'], level=1)
        tbl4 = doc.add_table(rows=1, cols=4)
        tbl4.style = "Table Grid"
        created_headers = [
            labels['quality.col.rule'],
            labels['quality.col.table'],
            labels['quality.col.dimension'],
            labels['quality.col.status'],
        ]
        for i, h in enumerate(created_headers):
            tbl4.rows[0].cells[i].text = h
            tbl4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for r in rules_created:
            row = tbl4.add_row().cells
            vals = [r["name"], r.get("table", "-"), r.get("dimension", "-"), r.get("status", "-")]
            for i, v in enumerate(vals):
                row[i].text = v
            # Tint the status column (col 3) when the value is a known status.
            _tint_cell(row[3], vals[3], _STATUS_CELL_COLOURS)

    # Recommendations
    recommendations = data.get("recommendations", [])
    if recommendations:
        doc.add_heading(labels['quality.recommendations'], level=1)
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")

    # Footer
    doc.add_paragraph("")
    footer_p = doc.add_paragraph(labels['quality.footer'])
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = GRAY

    ensure_output_dir(output_path)
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f"[OK] DOCX generated: {output_path} ({size_kb} KB)")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate data quality reports")
    parser.add_argument("--format", choices=["pdf", "docx", "md"], required=True,
                        help="Output format: pdf, docx or md")
    parser.add_argument("--output", required=True, help="Output file path")
    parser.add_argument("--input-json", dest="input_json", default=None,
                        help="JSON with the report data (string)")
    parser.add_argument("--input-file", dest="input_file", default=None,
                        help="Path to a JSON file with the report data")
    parser.add_argument("--lang", default=None,
                        help="Language code (e.g. en, es). Overrides `lang` in the "
                             "JSON input. Falls back to .agent_lang file, then 'en'.")
    parser.add_argument("--labels-json", dest="labels_json", default=None,
                        help='Label overrides as JSON string (e.g. '
                             "'{\"quality.executive_summary\":\"...\"}'). Highest priority.")
    parser.add_argument("--tone", default=None,
                        choices=[None, "default", "technical-minimal",
                                 "executive-editorial", "forensic"],
                        help="Visual tone for PDF and DOCX output. "
                             "Affects type pairing and accent palette. "
                             "See skill-guides/visual-craftsmanship.md.")
    args = parser.parse_args()

    try:
        data = load_input(args.input_json, args.input_file)
    except (json.JSONDecodeError, ValueError) as e:
        print(f"[ERROR] Could not load input data: {e}", file=sys.stderr)
        sys.exit(1)

    cli_label_overrides: dict | None = None
    if args.labels_json:
        try:
            cli_label_overrides = json.loads(args.labels_json)
        except json.JSONDecodeError as e:
            print(f"[ERROR] Invalid --labels-json: {e}", file=sys.stderr)
            sys.exit(1)

    labels = resolve_labels_for_data(
        data, cli_lang=args.lang, cli_labels=cli_label_overrides,
    )

    if args.format == "md":
        generate_markdown(data, args.output, labels=labels)
    elif args.format == "pdf":
        generate_pdf(data, args.output, labels=labels, tone=args.tone)
    elif args.format == "docx":
        generate_docx(data, args.output, labels=labels, tone=args.tone)


if __name__ == "__main__":
    main()

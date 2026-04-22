"""Generic, design-first DOCX builder for the docx-writer skill.

Exposes ``DOCXBuilder`` — a primitive-oriented API for producing .docx
documents with intentional design. Unlike the analytical ``DOCXGenerator``
inside data-analytics (which carries an opinionated scaffold with executive
summary / methodology / analysis / conclusions), this builder has no scaffold
baked in: callers compose documents out of primitives.

Typical use::

    from docx_builder import DOCXBuilder

    b = DOCXBuilder(page_size="A4", aesthetic_direction={"tone": "corporate"})
    b.add_cover(title="Policy Brief", subtitle="Retention and governance",
                author="Legal Team", metadata={"Ref": "POL-042"})
    b.add_heading("Scope", level=1)
    b.add_paragraph("This document defines the data retention policy ...")
    b.add_table(
        headers=["Dimension", "Target"],
        rows=[["Retention", "36 months"], ["Encryption", "AES-256"]],
        caption="Summary of commitments",
    )
    b.add_callout("All clauses require legal approval.", kind="warning")
    b.save("output/policy.docx")
"""
from __future__ import annotations

import base64
import re
import tempfile
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence, cast

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.styles.style import ParagraphStyle

# Load sibling `i18n.py` and `palette.py` under private module names. The
# monorepo ships three independent `i18n.py` files (data-analytics analyze,
# quality-report, and this skill) and two `palette.py` files; in a single
# Python process (e.g. a pytest session that imports several of them) a plain
# `from i18n import ...` resolves to whichever one entered `sys.modules['i18n']`
# first, which is undefined. Loading via `importlib.util` under a private name
# guarantees this skill always binds to its own catalogue.
import importlib.util as _imputil

_HERE = Path(__file__).resolve().parent


def _load_sibling(module_name: str, filename: str):
    spec = _imputil.spec_from_file_location(module_name, _HERE / filename)
    if spec is None or spec.loader is None:
        raise ImportError(f"cannot load sibling module {filename}")
    module = _imputil.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


_i18n_mod = _load_sibling("_docx_builder_i18n", "i18n.py")
_palette_mod = _load_sibling("_docx_builder_palette", "palette.py")

get_labels = _i18n_mod.get_labels
aesthetic_to_override_tokens = _palette_mod.aesthetic_to_override_tokens
get_palette = _palette_mod.get_palette

_IMAGE_WIDTH_INCHES_DEFAULT = 6.0
_CALLOUT_LIGHTEN = 0.78


def _lighten(rgb: tuple[int, int, int], factor: float = _CALLOUT_LIGHTEN) -> tuple[int, int, int]:
    return tuple(int(c + (255 - c) * factor) for c in rgb)  # type: ignore[return-value]


def _rgb_to_hex(rgb: tuple[int, int, int]) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_shading(cell, hex_fill: str) -> None:
    """Apply a CLEAR shading fill to a table cell (never SOLID — SOLID renders black)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:cantSplit"))
    if existing is None:
        existing = OxmlElement("w:cantSplit")
        tr_pr.append(existing)
    existing.set(qn("w:val"), "true")


def _set_row_repeat_as_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is None:
        existing = OxmlElement("w:tblHeader")
        tr_pr.append(existing)
    existing.set(qn("w:val"), "true")


def _set_cell_border(cell, *, top=None, bottom=None, left=None, right=None) -> None:
    """Set per-side borders on a cell. Pass ``None`` to remove that side."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side, spec in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = tc_borders.find(qn(f"w:{side}"))
        if spec is None:
            # remove that side if previously set
            if el is not None:
                tc_borders.remove(el)
            continue
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_borders.append(el)
        size, color, sty = spec
        el.set(qn("w:val"), sty or "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def _add_paragraph_border_bottom(paragraph, color_hex: str, size: int = 6) -> None:
    """Add a horizontal rule at the bottom of a paragraph (avoid tables-as-rules)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for existing in borders.findall(qn("w:bottom")):
        borders.remove(existing)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    borders.append(bottom)


def _add_page_number_field(run, placeholder: str = "1") -> None:
    """Insert a Word PAGE field into a run.

    Emits the full ``begin → instrText → separate → placeholder → end``
    sequence so LibreOffice, Word Online and Google Docs render a real
    number (not the raw instruction) even before the reader triggers a
    field update. The placeholder is the cached result; Word overwrites it
    on update, but LibreOffice uses it as-is on open.
    """
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = placeholder
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(placeholder_text)
    run._r.append(fld_char_end)


class DOCXBuilder:
    """Primitive-oriented DOCX builder.

    Parameters
    ----------
    page_size : str
        ``"A4"`` (default) or ``"Letter"``.
    orientation : str
        ``"portrait"`` (default) or ``"landscape"``.
    margin_cm : float
        Margin applied to all four sides. Default 2.5 cm (matches ISO).
    aesthetic_direction : dict | None
        Optional design-first dict with keys:

        - ``tone`` (str): one of ``"corporate"``, ``"academic"``, ``"modern"``.
          Unknown tones fall back to ``"corporate"``.
        - ``palette_override`` (dict): palette-key to hex or RGB overrides.
        - ``font_pair`` (list): ``[display, body]`` font names. ``body``
          becomes the document default; ``display`` is used for headings.

    author : str | None
        Document author written to core properties (and used in the cover
        page when no explicit ``author`` is passed to ``add_cover``).
    """

    def __init__(
        self,
        page_size: str = "A4",
        orientation: str = "portrait",
        margin_cm: float = 2.5,
        aesthetic_direction: dict | None = None,
        author: str | None = None,
    ) -> None:
        self._aesthetic = aesthetic_direction or {}
        tone = self._aesthetic.get("tone", "corporate")
        self._palette = get_palette(
            tone,
            override_tokens=aesthetic_to_override_tokens(self._aesthetic) or None,
        )
        self._author = author
        self._heading_font = self._resolve_heading_font()
        self._doc = self._build_document(page_size, orientation, margin_cm)
        self._apply_styles(self._doc)
        if author:
            self._doc.core_properties.author = author

    # ------------------------------------------------------------------
    # Setup
    # ------------------------------------------------------------------

    def _resolve_heading_font(self) -> str:
        fp = self._aesthetic.get("font_pair")
        if isinstance(fp, (list, tuple)) and len(fp) >= 1 and fp[0]:
            return str(fp[0])
        return str(self._palette.get("font_main", "Calibri"))

    def _build_document(self, page_size: str, orientation: str, margin_cm: float) -> DocumentObject:
        doc = Document()
        section = doc.sections[0]
        size = (page_size or "A4").upper()
        if size == "LETTER":
            w, h = Cm(21.59), Cm(27.94)
        else:
            w, h = Cm(21.0), Cm(29.7)
        if (orientation or "").lower() == "landscape":
            section.page_width, section.page_height = h, w
            from docx.enum.section import WD_ORIENT
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            section.page_width, section.page_height = w, h
        margin = Cm(margin_cm)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
        return doc

    def _apply_styles(self, doc: DocumentObject) -> None:
        p = self._palette
        font_main = str(p.get("font_main", "Calibri"))
        primary = p.get("primary", (26, 54, 93))
        primary_light = p.get("primary_light", (70, 130, 180))
        text_color = p.get("text", (33, 33, 33))
        text_muted = p.get("text_muted", (100, 100, 100))

        style_normal = cast(ParagraphStyle, doc.styles["Normal"])
        style_normal.font.name = font_main
        style_normal.font.size = Pt(11)
        if style_normal.font.color is not None:
            style_normal.font.color.rgb = RGBColor(*text_color)
        style_normal.paragraph_format.line_spacing = 1.15
        style_normal.paragraph_format.widow_control = True

        heading_spec = {
            1: (22, primary, 0),
            2: (16, primary, 1),
            3: (13, primary_light, 2),
        }
        for level, (size_pt, color, outline) in heading_spec.items():
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                style = cast(ParagraphStyle, doc.styles[style_name])
                style.font.name = self._heading_font
                style.font.size = Pt(size_pt)
                style.font.bold = True
                if style.font.color is not None:
                    style.font.color.rgb = RGBColor(*color)
                style.paragraph_format.keep_with_next = True
                # outlineLevel so automatic TOCs know the hierarchy
                p_pr = style.element.get_or_add_pPr()
                for existing in p_pr.findall(qn("w:outlineLvl")):
                    p_pr.remove(existing)
                ol = OxmlElement("w:outlineLvl")
                ol.set(qn("w:val"), str(outline))
                p_pr.append(ol)

        if "Caption" in doc.styles:
            cap_style = cast(ParagraphStyle, doc.styles["Caption"])
            cap_style.font.name = font_main
            cap_style.font.size = Pt(9)
            cap_style.font.italic = True
            if cap_style.font.color is not None:
                cap_style.font.color.rgb = RGBColor(*text_muted)

    # ------------------------------------------------------------------
    # Public primitives
    # ------------------------------------------------------------------

    @property
    def document(self) -> DocumentObject:
        """Access the underlying ``python-docx`` Document for advanced cases."""
        return self._doc

    def add_cover(
        self,
        title: str,
        subtitle: str | None = None,
        author: str | None = None,
        metadata: dict | None = None,
        lang: str | None = None,
        labels: dict[str, str] | None = None,
    ) -> "DOCXBuilder":
        """Add a cover page: title + subtitle + metadata block + page break."""
        resolved = get_labels(lang=lang, overrides=labels)
        primary = self._palette.get("primary", (26, 54, 93))
        text_muted = self._palette.get("text_muted", (100, 100, 100))
        font_main = str(self._palette.get("font_main", "Calibri"))

        for _ in range(6):
            self._doc.add_paragraph("")

        title_para = self._doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title_para.add_run(title)
        tr.font.size = Pt(32)
        tr.font.bold = True
        tr.font.color.rgb = RGBColor(*primary)
        tr.font.name = self._heading_font

        if subtitle:
            sub_para = self._doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sub_para.add_run(subtitle)
            sr.font.size = Pt(16)
            sr.font.color.rgb = RGBColor(*text_muted)
            sr.font.name = font_main

        self._doc.add_paragraph("")

        meta_lines: list[str] = []
        effective_author = author or self._author
        if effective_author:
            meta_lines.append(f"{resolved['cover.author']}: {effective_author}")
        if metadata:
            for k, v in metadata.items():
                meta_lines.append(f"{k}: {v}")
        meta_lines.append(f"{resolved['cover.date']}: {date.today().strftime('%d/%m/%Y')}")

        for line in meta_lines:
            mp = self._doc.add_paragraph()
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mr = mp.add_run(line)
            mr.font.size = Pt(12)
            mr.font.color.rgb = RGBColor(*text_muted)
            mr.font.name = font_main

        self._doc.add_page_break()
        return self

    def add_heading(
        self,
        text: str,
        level: int = 1,
        page_break_before: bool = False,
    ) -> "DOCXBuilder":
        level = max(1, min(6, int(level)))
        style_name = f"Heading {level}"
        para = self._doc.add_paragraph(text, style=style_name)
        # Headings should stay with the next paragraph; this protects
        # against the heading ending up alone at the bottom of a page.
        para.paragraph_format.keep_with_next = True
        if page_break_before:
            para.paragraph_format.page_break_before = True
        return self

    def add_paragraph(self, text: str) -> "DOCXBuilder":
        """Add a paragraph. Inline bold / italic expressed as ``**text**`` / ``*text*``."""
        para = self._doc.add_paragraph()
        self._add_formatted_runs(para, text)
        return self

    # **bold** (non-greedy so it can contain single *...* italics),
    # *italic* (lookbehind prevents matching inside bold),
    # `code` (monospaced).
    _INLINE_PATTERN = re.compile(
        r"\*\*.+?\*\*"
        r"|(?<!\*)\*[^*]+\*"
        r"|`[^`]+`"
    )
    # Inside a bold block: italics and code markers still apply.
    _BOLD_INNER_PATTERN = re.compile(r"\*([^*]+)\*|`([^`]+)`")

    def _add_formatted_runs(self, para, text: str) -> None:
        """Minimal inline formatting: **bold**, *italic*, `code`.

        Italics and inline code nested inside bold (``**bold with *italic* inside**``)
        are honoured and rendered as a single paragraph with bold-italic or
        bold-monospaced runs at the right positions.
        """
        pos = 0
        for m in self._INLINE_PATTERN.finditer(text):
            if m.start() > pos:
                para.add_run(text[pos : m.start()])
            chunk = m.group(0)
            if chunk.startswith("**"):
                self._emit_bold_block(para, chunk[2:-2])
            elif chunk.startswith("*"):
                r = para.add_run(chunk[1:-1])
                r.italic = True
            else:  # `code`
                r = para.add_run(chunk[1:-1])
                r.font.name = str(self._palette.get("font_mono", "Consolas"))
            pos = m.end()
        if pos < len(text):
            para.add_run(text[pos:])

    def _emit_bold_block(self, para, inner: str) -> None:
        """Emit runs for a bold block, recognising nested italic and code markers."""
        mono = str(self._palette.get("font_mono", "Consolas"))
        pos = 0
        for mi in self._BOLD_INNER_PATTERN.finditer(inner):
            if mi.start() > pos:
                r = para.add_run(inner[pos : mi.start()])
                r.bold = True
            if mi.group(0).startswith("*"):
                r = para.add_run(mi.group(1))
                r.bold = True
                r.italic = True
            else:  # `code`
                r = para.add_run(mi.group(2))
                r.bold = True
                r.font.name = mono
            pos = mi.end()
        if pos < len(inner):
            r = para.add_run(inner[pos:])
            r.bold = True

    def add_figure(
        self,
        path: str | Path,
        caption: str | None = None,
        width_inches: float = _IMAGE_WIDTH_INCHES_DEFAULT,
    ) -> "DOCXBuilder":
        p = Path(path)
        if p.is_file():
            self._doc.add_picture(str(p), width=Inches(width_inches))
            img_para = self._doc.paragraphs[-1]
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.keep_with_next = True
            img_para.paragraph_format.keep_together = True
        else:
            missing = self._doc.add_paragraph(f"[Image not found: {p}]")
            missing.runs[0].font.italic = True
            missing.runs[0].font.color.rgb = RGBColor(150, 150, 150)
            missing.paragraph_format.keep_with_next = True

        if caption:
            cap_para = self._doc.add_paragraph(caption, style="Caption")
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._doc.add_paragraph("")
        return self

    def add_table(
        self,
        headers: Sequence[str],
        rows: Sequence[Sequence[str]],
        caption: str | None = None,
        style: str = "shaded-header",
        column_widths_cm: Sequence[float] | None = None,
        cell_colors: dict[tuple[int, int], tuple[int, int, int]] | None = None,
    ) -> "DOCXBuilder":
        """Add a table with deliberate styling.

        ``style`` controls visual treatment:
        - ``"shaded-header"`` (default) — filled header row with primary fill and
          white text; subtle rule below the header and between rows.
        - ``"minimal"`` — no fill, only a thin rule below the header.

        ``cell_colors`` is an optional dict mapping ``(row_index, col_index)``
        (both 1-indexed starting with the first body row; the header row is
        row 0 and is never targetable from this map) to an ``(R, G, B)`` tuple
        applied as the text colour of that cell. Useful for status / priority
        columns that should stand out (e.g., OK green, KO red).
        """
        if not headers:
            return self
        primary = self._palette.get("primary", (26, 54, 93))
        bg_alt = self._palette.get("bg_alt", (247, 250, 252))
        border = self._palette.get("border", (226, 232, 240))
        text_muted = self._palette.get("text_muted", (100, 100, 100))

        table = self._doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        if column_widths_cm:
            if len(column_widths_cm) != len(headers):
                raise ValueError("column_widths_cm length must match headers")
            for row in table.rows:
                for idx, w in enumerate(column_widths_cm):
                    row.cells[idx].width = Cm(float(w))

        # Header
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(h))
            run.font.bold = True
            run.font.size = Pt(10)
            if style == "shaded-header":
                _set_cell_shading(cell, _rgb_to_hex(primary))
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                run.font.color.rgb = RGBColor(*primary)
                _set_cell_border(
                    cell,
                    bottom=(6, _rgb_to_hex(primary), "single"),
                )

        _set_row_repeat_as_header(table.rows[0])

        # Body
        for r_idx, row_data in enumerate(rows, start=1):
            for c_idx, value in enumerate(row_data):
                cell = table.rows[r_idx].cells[c_idx]
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(str(value))
                run.font.size = Pt(10)
                if cell_colors is not None:
                    colour = cell_colors.get((r_idx, c_idx))
                    if colour is not None:
                        run.font.color.rgb = RGBColor(*colour)
                        run.font.bold = True
                if style == "shaded-header" and r_idx % 2 == 0:
                    _set_cell_shading(cell, _rgb_to_hex(bg_alt))
                if style == "minimal":
                    _set_cell_border(
                        cell,
                        bottom=(4, _rgb_to_hex(border), "single"),
                    )

        # Avoid splitting any single row across a page break
        for row in table.rows:
            _set_row_cant_split(row)

        if caption:
            cap_para = self._doc.add_paragraph(caption, style="Caption")
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap_para.runs:
                run.font.color.rgb = RGBColor(*text_muted)

        self._doc.add_paragraph("")
        return self

    def add_callout(self, text: str, kind: str = "info") -> "DOCXBuilder":
        """Add an info/success/warning/danger callout as a shaded one-cell table."""
        palette_key = {
            "info": "primary_light",
            "success": "success",
            "warning": "warning",
            "danger": "danger",
        }.get(kind, "primary_light")
        base = self._palette.get(palette_key, (49, 130, 206))
        fill = _lighten(base, 0.85)
        stripe = base

        table = self._doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _set_cell_shading(cell, _rgb_to_hex(fill))
        _set_cell_border(
            cell,
            left=(24, _rgb_to_hex(stripe), "single"),
            top=None,
            bottom=None,
            right=None,
        )
        para = cell.paragraphs[0]
        self._add_formatted_runs(para, text)
        for run in para.runs:
            run.font.size = Pt(10)
        _set_row_cant_split(table.rows[0])
        self._doc.add_paragraph("")
        return self

    def add_list(self, items: Iterable[str], ordered: bool = False) -> "DOCXBuilder":
        style = "List Number" if ordered else "List Bullet"
        for item in items:
            self._doc.add_paragraph(str(item), style=style)
        return self

    def add_code_block(self, text: str, language: str | None = None) -> "DOCXBuilder":
        """Render a monospaced block with subtle background shading.

        ``language`` is accepted for API parity with code-fence consumers but
        is not rendered (DOCX has no native syntax-highlighting primitive).
        """
        _ = language  # noqa: F841 — reserved for future highlighters
        bg_alt = self._palette.get("bg_alt", (247, 250, 252))
        text_color = self._palette.get("text", (33, 33, 33))
        font_mono = str(self._palette.get("font_mono", "Consolas"))

        table = self._doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _set_cell_shading(cell, _rgb_to_hex(bg_alt))
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for idx, line in enumerate(text.splitlines() or [text]):
            if idx > 0:
                para.add_run().add_break()
            run = para.add_run(line)
            run.font.name = font_mono
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(*text_color)
        _set_row_cant_split(table.rows[0])
        self._doc.add_paragraph("")
        return self

    def add_page_break(self) -> "DOCXBuilder":
        para = self._doc.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)
        return self

    def add_horizontal_rule(self) -> "DOCXBuilder":
        """Add a horizontal rule as a paragraph border — never a 1-cell table."""
        border_rgb = self._palette.get("border", (226, 232, 240))
        para = self._doc.add_paragraph()
        _add_paragraph_border_bottom(para, _rgb_to_hex(border_rgb), size=6)
        return self

    def add_markdown_block(self, md_text: str) -> "DOCXBuilder":
        """Render a markdown string as DOCX primitives.

        Converts ``md_text`` to HTML via ``markdown.markdown`` and delegates to
        :meth:`add_html_block`, so it gains everything HTML supports, including
        base64-encoded inline images (``![alt](data:image/png;base64,...)``).
        """
        import markdown as md_lib

        html = md_lib.markdown(
            md_text, extensions=["tables", "fenced_code"]
        )
        return self.add_html_block(html)

    def add_html_block(self, html_text: str) -> "DOCXBuilder":
        """Render an HTML string as DOCX primitives.

        Supported tags:
        - ``<h1>``…``<h6>`` — headings at the corresponding level.
        - ``<p>`` — paragraph with inline ``<strong>``/``<b>``, ``<em>``/``<i>``,
          ``<code>`` preserved.
        - ``<ul>``/``<ol>`` — bullet / numbered lists.
        - ``<pre>``/``<code>`` — fenced code blocks.
        - ``<hr>`` — horizontal rule (paragraph border).
        - ``<table>`` — table (uses ``shaded-header`` style).
        - ``<figure>`` / ``<img>`` — figures. ``src`` may be a normal path or a
          ``data:image/...;base64,...`` URI; base64 is decoded to a temporary
          PNG and embedded via :meth:`add_figure`.
        - ``<div class="callout [info|success|warning|danger]">`` and
          ``<blockquote>`` — callouts.

        Unknown tags fall back to rendering their text content as a paragraph.
        """
        from bs4 import BeautifulSoup, NavigableString, Tag

        soup = BeautifulSoup(html_text, "html.parser")

        for node in soup.children:
            if isinstance(node, NavigableString):
                text = str(node).strip()
                if text:
                    self.add_paragraph(text)
                continue
            if not isinstance(node, Tag):
                continue
            tag = (node.name or "").lower()

            if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                self.add_heading(node.get_text().strip(), level=int(tag[1]))

            elif tag == "p":
                inner = self._html_inline_to_md(node)
                if inner.strip():
                    self.add_paragraph(inner)

            elif tag in {"ul", "ol"}:
                items = [
                    self._html_inline_to_md(li)
                    for li in node.find_all("li", recursive=False)
                ]
                self.add_list([i for i in items if i.strip()], ordered=(tag == "ol"))

            elif tag == "hr":
                self.add_horizontal_rule()

            elif tag == "pre":
                code = node.find("code")
                text = code.get_text() if code else node.get_text()
                self.add_code_block(text.rstrip("\n"))

            elif tag == "table":
                headers = [th.get_text().strip() for th in node.select("thead th")]
                if not headers:
                    # Fallback: first row as header when no <thead> is present
                    first = node.find("tr")
                    if first:
                        headers = [c.get_text().strip() for c in first.find_all(("th", "td"))]
                rows = []
                body_rows = node.select("tbody tr") or node.find_all("tr")[1:]
                for tr in body_rows:
                    rows.append([td.get_text().strip() for td in tr.find_all(("td", "th"))])
                if headers and rows:
                    self.add_table(headers, rows)

            elif tag == "figure":
                img = node.find("img")
                cap_tag = node.find("figcaption")
                caption = cap_tag.get_text(strip=True) if cap_tag else ""
                if img is not None:
                    self._render_img_tag(img, caption_fallback=caption)

            elif tag == "img":
                self._render_img_tag(node)

            elif tag == "div":
                classes = node.get("class") or []
                if "callout" in classes:
                    kind = next(
                        (c for c in classes if c in {"info", "success", "warning", "danger"}),
                        "info",
                    )
                    self.add_callout(node.get_text(strip=True), kind=kind)
                else:
                    # Recurse into div content
                    self.add_html_block(node.decode_contents())

            elif tag == "blockquote":
                self.add_callout(node.get_text(strip=True), kind="info")

            elif tag == "br":
                self._doc.add_paragraph("")

            else:
                # Fallback for any other tag with text content
                text = node.get_text(strip=True)
                if text:
                    self.add_paragraph(self._html_inline_to_md(node))

        return self

    def _render_img_tag(self, img_tag, caption_fallback: str = "") -> None:
        """Insert an <img> — resolve `data:` URIs to temporary PNG files."""
        src = img_tag.get("src", "") or ""
        alt = img_tag.get("alt", "") or ""
        caption = caption_fallback or alt
        if src.startswith("data:"):
            tmp = self._decode_data_uri(src)
            try:
                self.add_figure(tmp, caption=caption or None)
            finally:
                Path(tmp).unlink(missing_ok=True)
        elif src:
            self.add_figure(src, caption=caption or None)

    @staticmethod
    def _decode_data_uri(data_uri: str) -> str:
        """Decode a ``data:image/...;base64,...`` URI to a temporary PNG file."""
        payload = data_uri.split(",", 1)[1] if "," in data_uri else data_uri
        img_bytes = base64.b64decode(payload)
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        try:
            tmp.write(img_bytes)
        finally:
            tmp.close()
        return tmp.name

    @staticmethod
    def _html_inline_to_md(node) -> str:
        """Flatten a tag's children to lightweight markdown so _add_formatted_runs
        can pick up bold / italic / code."""
        from bs4 import NavigableString, Tag

        parts: list[str] = []
        for child in getattr(node, "children", []):
            if isinstance(child, NavigableString):
                parts.append(str(child))
            elif isinstance(child, Tag):
                name = (child.name or "").lower()
                text = DOCXBuilder._html_inline_to_md(child)
                if name in {"strong", "b"}:
                    parts.append(f"**{text}**")
                elif name in {"em", "i"}:
                    parts.append(f"*{text}*")
                elif name == "code":
                    parts.append(f"`{text}`")
                elif name == "br":
                    parts.append("\n")
                else:
                    parts.append(text)
        return "".join(parts)

    # ------------------------------------------------------------------
    # Header / footer helpers
    # ------------------------------------------------------------------

    def set_footer_page_numbers(self, lang: str | None = None) -> "DOCXBuilder":
        """Configure a centred "Page X" footer on the default section."""
        labels = get_labels(lang=lang)
        font_main = str(self._palette.get("font_main", "Calibri"))
        text_muted = self._palette.get("text_muted", (100, 100, 100))
        section = self._doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()
        run_label = para.add_run(f"{labels['page.footer_page_label']} ")
        run_label.font.name = font_main
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = RGBColor(*text_muted)
        run_num = para.add_run()
        run_num.font.name = font_main
        run_num.font.size = Pt(9)
        run_num.font.color.rgb = RGBColor(*text_muted)
        _add_page_number_field(run_num)
        return self

    # ------------------------------------------------------------------
    # Output
    # ------------------------------------------------------------------

    def save(self, output_path: str | Path) -> Path:
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(path))
        return path

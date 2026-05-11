#!/usr/bin/env python3
"""Quick extraction from a Word document.

Returns structured Markdown on stdout. Diagnostics go to stderr. Exit code 0
on success, 1 on failure. The extractor chain is: pandoc -> python-docx ->
raw-XML walk. The first one that produces non-empty text wins.

Usage:
    python3 quick_extract.py document.docx
    python3 quick_extract.py document.docx --no-tables
    python3 quick_extract.py document.docx --tool pandoc
    cat document.docx | python3 quick_extract.py -
"""
from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

SUPPORTED_TOOLS = ("auto", "pandoc", "python-docx", "unzip")


def _log(msg: str) -> None:
    print(msg, file=sys.stderr)


def _read_stdin_to_tempfile() -> Path:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(sys.stdin.buffer.read())
    tmp.close()
    return Path(tmp.name)


def _convert_legacy_doc(path: Path) -> Path:
    """Convert a legacy binary .doc to .docx via soffice. Returns new path."""
    if shutil.which("soffice") is None:
        raise RuntimeError(
            "File looks like a legacy .doc but 'soffice' (LibreOffice) is not "
            "installed. Install libreoffice or libreoffice-writer."
        )
    out_dir = Path(tempfile.mkdtemp(prefix="docx_conv_"))
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "docx",
         "--outdir", str(out_dir), str(path)],
        check=True, capture_output=True,
    )
    converted = out_dir / (path.stem + ".docx")
    if not converted.exists():
        raise RuntimeError(f"LibreOffice produced no output for {path}")
    return converted


def _is_docx_zip(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path) as z:
            return "word/document.xml" in z.namelist()
    except zipfile.BadZipFile:
        return False


def _extract_with_pandoc(path: Path, include_tables: bool) -> str:
    if shutil.which("pandoc") is None:
        raise RuntimeError("pandoc not installed")
    result = subprocess.run(
        ["pandoc", "--from", "docx", "--to", "gfm", str(path)],
        capture_output=True, text=True, check=False,
    )
    if result.returncode != 0:
        raise RuntimeError(f"pandoc failed: {result.stderr.strip()}")
    text = result.stdout
    if not include_tables:
        text = _strip_markdown_tables(text)
    return text.strip()


def _strip_markdown_tables(md: str) -> str:
    """Remove GFM table blocks from markdown output."""
    lines = md.splitlines()
    out = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            continue
        if in_table and not stripped:
            in_table = False
            out.append("")
            continue
        if in_table:
            continue
        out.append(line)
    return "\n".join(out)


def _extract_with_python_docx(path: Path, include_tables: bool) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx not installed") from exc

    doc = Document(str(path))
    parts: list[str] = []

    title = None
    if doc.core_properties and doc.core_properties.title:
        title = doc.core_properties.title.strip() or None
    if title:
        parts.append(f"# {title}\n")

    # Iterate paragraphs and tables in document order by walking the body xml.
    body = doc.element.body
    from docx.oxml.ns import qn

    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            para = [p for p in doc.paragraphs if p._p is child]
            if not para:
                continue
            p = para[0]
            text = p.text.strip()
            if not text:
                continue
            style = (p.style.name or "").lower() if p.style else ""
            if style.startswith("heading "):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 2
                parts.append(f"{'#' * max(1, min(level, 6))} {text}")
            else:
                parts.append(text)
        elif child.tag == qn("w:tbl") and include_tables:
            table = [t for t in doc.tables if t._tbl is child]
            if not table:
                continue
            parts.append(_table_to_markdown(table[0]))

    return "\n\n".join(parts).strip()


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        seen: list[int] = []
        cells_text: list[str] = []
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.append(cid)
            cells_text.append(cell.text.strip().replace("|", "\\|") or " ")
        rows.append(cells_text)

    if not rows:
        return ""
    max_cols = max(len(r) for r in rows)
    rows = [r + [" "] * (max_cols - len(r)) for r in rows]

    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * max_cols) + " |"
    body = ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return "\n".join([header, sep, *body])


def _extract_with_unzip(path: Path, include_tables: bool) -> str:
    from lxml import etree  # type: ignore[import-untyped]

    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": NS_W}

    with zipfile.ZipFile(path) as z:
        with z.open("word/document.xml") as f:
            tree = etree.parse(f)

    parts: list[str] = []
    body = tree.getroot().find("w:body", ns)
    if body is None:
        return ""

    for child in body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = "".join(t.text or "" for t in child.iter(f"{{{NS_W}}}t")).strip()
            if text:
                parts.append(text)
        elif tag == "tbl" and include_tables:
            rows = []
            for row in child.findall("w:tr", ns):
                cells = []
                for cell in row.findall("w:tc", ns):
                    cell_text = "".join(
                        t.text or "" for t in cell.iter(f"{{{NS_W}}}t")
                    ).strip().replace("|", "\\|") or " "
                    cells.append(cell_text)
                rows.append(cells)
            if rows:
                max_cols = max(len(r) for r in rows)
                rows = [r + [" "] * (max_cols - len(r)) for r in rows]
                header = "| " + " | ".join(rows[0]) + " |"
                sep = "| " + " | ".join(["---"] * max_cols) + " |"
                body_lines = ["| " + " | ".join(r) + " |" for r in rows[1:]]
                parts.append("\n".join([header, sep, *body_lines]))
    return "\n\n".join(parts).strip()


def extract(path: Path, tool: str = "auto", include_tables: bool = True) -> str:
    chain = ["pandoc", "python-docx", "unzip"] if tool == "auto" else [tool]
    errors: list[str] = []
    for name in chain:
        try:
            if name == "pandoc":
                text = _extract_with_pandoc(path, include_tables)
            elif name == "python-docx":
                text = _extract_with_python_docx(path, include_tables)
            elif name == "unzip":
                text = _extract_with_unzip(path, include_tables)
            else:
                raise RuntimeError(f"unknown tool '{name}'")
        except Exception as exc:  # noqa: BLE001
            _log(f"[{name}] failed: {exc}")
            errors.append(f"{name}: {exc}")
            continue
        if text.strip():
            _log(f"[{name}] produced {len(text)} chars")
            return text
        _log(f"[{name}] produced empty output, trying next tool")
    raise RuntimeError(
        "All extractors failed or produced empty output. Errors: "
        + "; ".join(errors)
    )


def main() -> int:
    parser = argparse.ArgumentParser(
        description=(__doc__ or "").split("\n\n")[0]
    )
    parser.add_argument(
        "path",
        help="Path to the .docx file, or '-' to read from stdin.",
    )
    parser.add_argument("--no-tables", action="store_true",
                        help="Skip tables, emit only prose.")
    parser.add_argument("--tool", choices=SUPPORTED_TOOLS, default="auto",
                        help="Force a specific extractor (default: auto fallback).")
    args = parser.parse_args()

    if args.path == "-":
        path = _read_stdin_to_tempfile()
    else:
        path = Path(args.path)
        if not path.exists():
            _log(f"error: file not found: {path}")
            return 1

    # Detect legacy .doc and convert up front
    try:
        if not _is_docx_zip(path):
            _log(f"input is not a .docx ZIP; attempting legacy .doc conversion")
            path = _convert_legacy_doc(path)
    except Exception as exc:  # noqa: BLE001
        _log(f"error: cannot read input: {exc}")
        return 1

    try:
        text = extract(path, tool=args.tool, include_tables=not args.no_tables)
    except Exception as exc:  # noqa: BLE001
        _log(f"error: {exc}")
        return 1

    sys.stdout.write(text)
    sys.stdout.write("\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
